"""
policy_generator.py
Waymark HR Group LLC

Generates a professional HR Policy Handbook as a Word document (.docx).
Saves the document to docs/reports/.

Usage:
    python policy_generator.py "Company Name" "Industry" "150" "New York"

Arguments (all required):
    1. Company name
    2. Industry
    3. Company size (number of employees)
    4. State (for legal compliance references)
"""

import sys
import argparse
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

CONSULTANT_NAME    = "Waymark HR Group LLC"
CONSULTANT_ADDRESS = "Buffalo, NY"
CONSULTANT_PHONE   = "716-225-6347"
CONSULTANT_EMAIL   = "Jamie.Wahler@waymarkhrgroup.com"
CONSULTANT_WEBSITE = "www.waymarkhrgroup.com"

COLOR_NAVY       = RGBColor(0x1A, 0x3A, 0x5C)
COLOR_TEAL       = RGBColor(0x00, 0x7A, 0x87)
COLOR_LIGHT_GRAY = RGBColor(0xF4, 0xF6, 0xF8)
COLOR_TEXT       = RGBColor(0x2C, 0x2C, 0x2C)
COLOR_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_MUTED      = RGBColor(0x88, 0x88, 0x88)

REPORTS_DIR = Path(__file__).resolve().parents[2] / "docs" / "reports"
LOGO_PATH   = Path(__file__).resolve().parents[2] / "docs" / "templates" / "Waymark_HR_Group_Logo_Full.png"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_paragraph(doc, text="", bold=False, italic=False, font_size=11,
                  color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        run.font.color.rgb = color or COLOR_TEXT
    return p


def add_section_heading(doc, text):
    """Teal underlined section heading for a policy section label (e.g. PURPOSE)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = COLOR_TEAL
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "007A87")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_policy_title(doc, number, title):
    """Navy bold policy title with number."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"Policy {number}: {title}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_NAVY
    return p


def add_body(doc, text, space_after=6):
    return add_paragraph(doc, text, font_size=10.5, space_after=space_after)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_TEXT
    return p


def add_page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document header & footer
# ---------------------------------------------------------------------------

def build_header(doc, company_name: str):
    """Logo + navy branded title block."""
    if LOGO_PATH.exists():
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.paragraph_format.space_before = Pt(0)
        logo_para.paragraph_format.space_after = Pt(8)
        logo_para.add_run().add_picture(str(LOGO_PATH), width=Inches(3.0))

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1A3A5C")

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(12)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run("HR POLICY HANDBOOK")
    r1.bold = True
    r1.font.size = Pt(16)
    r1.font.color.rgb = COLOR_WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(company_name)
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0xB0, 0xC8, 0xD8)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(12)
    r3 = p3.add_run(f"Effective Date: {datetime.now().strftime('%B %d, %Y')}  |  Prepared by {CONSULTANT_NAME}")
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0xC8, 0xD8, 0xE8)

    doc.add_paragraph()


def build_document_footer(doc):
    """Add footer with consultant contact info to all pages."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    ft = footer.paragraphs[0]
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ft.paragraph_format.space_before = Pt(4)

    # Top border on footer
    pPr = ft._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "007A87")
    pBdr.append(top)
    pPr.append(pBdr)

    run = ft.add_run(
        f"{CONSULTANT_NAME}  |  {CONSULTANT_PHONE}  |  {CONSULTANT_EMAIL}  |  {CONSULTANT_WEBSITE}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED

    # Page number on second line
    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    r2 = p2.add_run("Page ")
    r2.font.size = Pt(8)
    r2.font.color.rgb = COLOR_MUTED

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = p2.add_run()
    run2.font.size = Pt(8)
    run2.font.color.rgb = COLOR_MUTED
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)


def build_intro(doc, company_name, industry, company_size, state):
    add_section_heading(doc, "Introduction")
    add_body(doc,
        f"This HR Policy Handbook has been prepared for {company_name}, a {industry} organization "
        f"with approximately {company_size} employees operating in the state of {state}. "
        f"It establishes the standards, expectations, and guidelines that govern the employment "
        f"relationship and workplace conduct for all employees."
    )
    add_body(doc,
        "All employees are required to read, understand, and comply with the policies contained "
        "in this handbook. This document does not constitute an employment contract. The company "
        "reserves the right to amend, modify, or rescind any policy at any time, with or without "
        "advance notice."
    )
    add_body(doc,
        f"Employees with questions regarding any policy should contact their direct manager or "
        f"the Human Resources department. Legal questions should be directed to qualified employment "
        f"counsel familiar with {state} employment law.",
        space_after=14
    )


# ---------------------------------------------------------------------------
# Policy sections
# ---------------------------------------------------------------------------

def build_policy(doc, number, title, purpose, scope, details, responsibilities, consequences):
    """Render a single policy with all five standard sections."""
    add_page_break(doc)
    add_policy_title(doc, number, title)

    add_section_heading(doc, "Purpose")
    add_body(doc, purpose)

    add_section_heading(doc, "Scope")
    add_body(doc, scope)

    add_section_heading(doc, "Policy")
    for item in details:
        if item.startswith("__bullet__"):
            add_bullet(doc, item[len("__bullet__"):])
        else:
            add_body(doc, item)

    add_section_heading(doc, "Employee Responsibilities")
    for item in responsibilities:
        add_bullet(doc, item)

    add_section_heading(doc, "Consequences for Violations")
    add_body(doc, consequences)


def policy_eeo(company_name, state):
    return dict(
        title="Equal Employment Opportunity (EEO)",
        purpose=(
            f"{company_name} is an equal opportunity employer committed to providing a workplace "
            f"free from discrimination. We believe that a diverse and inclusive workforce strengthens "
            f"our organization and the communities we serve."
        ),
        scope=(
            f"This policy applies to all employees, applicants, contractors, and vendors of "
            f"{company_name} in all locations, including remote workers."
        ),
        details=[
            f"{company_name} does not discriminate in any aspect of employment — including hiring, "
            f"promotion, compensation, training, discipline, or termination — on the basis of race, "
            f"color, religion, sex, national origin, age, disability, genetic information, veteran "
            f"status, sexual orientation, gender identity, or any other characteristic protected "
            f"under federal, {state} state, or applicable local law.",
            "This commitment applies to all employment practices and is upheld by every level of "
            "management. Reasonable accommodations will be provided to qualified individuals with "
            "disabilities or sincerely held religious beliefs, unless doing so would create an undue "
            "hardship for the company.",
            "All hiring decisions are based solely on job-related qualifications, merit, and "
            "business needs."
        ],
        responsibilities=[
            "Treat all colleagues, applicants, and business partners with dignity and respect.",
            "Report any observed or experienced discriminatory conduct to HR immediately.",
            "Cooperate fully with any investigation related to an EEO complaint.",
            "Refrain from retaliating against anyone who reports a concern in good faith.",
        ],
        consequences=(
            "Violations of this policy will result in disciplinary action, up to and including "
            "termination of employment. The company will also take appropriate action against "
            "any third party found to have engaged in discriminatory conduct."
        ),
    )


def policy_anti_harassment(company_name, state):
    return dict(
        title="Anti-Harassment & Anti-Discrimination",
        purpose=(
            f"{company_name} is committed to maintaining a professional work environment in which "
            f"all individuals are treated with respect. Harassment and discrimination of any kind "
            f"are strictly prohibited."
        ),
        scope=(
            "This policy applies to all employees, contractors, clients, and visitors. It covers "
            "conduct occurring in the workplace, at company-sponsored events, and through company "
            "communication systems, including email and messaging platforms."
        ),
        details=[
            "Harassment includes any unwelcome verbal, visual, or physical conduct based on a "
            "protected characteristic that creates an intimidating, hostile, or offensive work "
            "environment, or that interferes with an individual's work performance.",
            "Sexual harassment includes unwelcome sexual advances, requests for sexual favors, "
            "and other verbal or physical conduct of a sexual nature when submission to such "
            "conduct is made a condition of employment, or when such conduct unreasonably "
            "interferes with an individual's work performance.",
            f"{state} state law may provide additional protections beyond federal requirements. "
            f"The company complies with all applicable state and local anti-harassment statutes.",
            "Retaliation against any person who reports harassment or participates in an "
            "investigation is strictly prohibited and will itself be treated as a violation "
            "of this policy.",
        ],
        responsibilities=[
            "Maintain professional, respectful conduct toward all colleagues and third parties.",
            "Report any witnessed or experienced harassment promptly to HR or a supervisor.",
            "Cooperate honestly and fully in any investigation.",
            "Maintain confidentiality regarding all complaint and investigation details.",
        ],
        consequences=(
            "Any employee found to have engaged in harassment or discrimination will be subject "
            "to disciplinary action up to and including immediate termination. The company may "
            "also notify appropriate authorities where conduct may constitute a violation of law."
        ),
    )


def policy_attendance(company_name, state):
    return dict(
        title="Attendance & Punctuality",
        purpose=(
            f"Consistent attendance and punctuality are essential to the success of {company_name} "
            f"and to maintaining a productive, fair work environment for all team members."
        ),
        scope=(
            "This policy applies to all full-time and part-time employees. Specific scheduling "
            "requirements may vary by department and will be communicated by the employee's manager."
        ),
        details=[
            f"Employees are expected to report to work on time and fully prepared at the start of "
            f"each scheduled shift. Core business hours are established by department management "
            f"in accordance with {company_name} operational needs.",
            "If an employee is unable to report to work, they must notify their manager as early "
            "as possible and no later than the start of their scheduled shift. Failure to provide "
            "timely notice without valid justification is considered an unexcused absence.",
            "Excessive absenteeism or tardiness — whether excused or unexcused — may negatively "
            "impact performance evaluations and may result in disciplinary action.",
            f"{state} state law requirements regarding leave, including any applicable sick leave "
            f"mandates, are incorporated into this policy and will be honored in all circumstances.",
        ],
        responsibilities=[
            "Arrive on time and ready to work for every scheduled shift.",
            "Notify the manager promptly when absence or tardiness is anticipated.",
            "Use the appropriate system or process to record absences and time off.",
            "Obtain approval for planned absences in advance whenever possible.",
        ],
        consequences=(
            "Unexcused absences or chronic tardiness will be addressed through a progressive "
            "disciplinary process, which may include a verbal warning, written warning, final "
            "written warning, and termination of employment, depending on severity and frequency."
        ),
    )


def policy_pto(company_name, state):
    return dict(
        title="Paid Time Off (PTO) & Leave",
        purpose=(
            f"{company_name} recognizes that employees need time away from work for rest, "
            f"personal matters, illness, and family obligations. This policy outlines the "
            f"company's approach to paid time off and leave."
        ),
        scope=(
            "This policy applies to all regular full-time employees. Part-time employees may "
            "be eligible for prorated benefits as outlined in their offer letter or employment "
            "agreement."
        ),
        details=[
            "PTO accrues from the first day of employment based on the company's published "
            "accrual schedule. Employees are encouraged to use PTO throughout the year.",
            "PTO requests must be submitted to the employee's manager for approval. Requests "
            "should be made as far in advance as possible — a minimum of two weeks' notice is "
            "expected for planned absences of three or more consecutive days.",
            "Unused PTO may be carried over subject to the annual carryover cap established by "
            "the company. PTO is not paid out upon resignation or termination unless required "
            f"by {state} state law.",
            f"In addition to PTO, {company_name} provides leave in accordance with applicable "
            f"federal and {state} state law, including but not limited to: Family and Medical "
            f"Leave Act (FMLA), military leave (USERRA), jury duty leave, bereavement leave, "
            f"and any {state}-specific leave entitlements.",
        ],
        responsibilities=[
            "Submit PTO requests through the designated HR system in a timely manner.",
            "Communicate leave needs clearly and early to minimize operational disruption.",
            "Provide appropriate documentation when required (e.g., FMLA certification, jury summons).",
            "Return to work as scheduled following any approved leave.",
        ],
        consequences=(
            "Misrepresentation of the reason for leave, or failure to return from an approved "
            "leave without prior authorization, may result in disciplinary action up to and "
            "including termination. Abuse of PTO or leave policies will be addressed accordingly."
        ),
    )


def policy_code_of_conduct(company_name):
    return dict(
        title="Code of Conduct",
        purpose=(
            f"The {company_name} Code of Conduct establishes the behavioral standards expected "
            f"of all employees. Upholding these standards is essential to maintaining a respectful, "
            f"ethical, and productive workplace."
        ),
        scope=(
            "This policy applies to all employees at all levels of the organization, including "
            "part-time, temporary, and contract workers. It applies to conduct in the workplace, "
            "at company events, and in any context where the employee represents the company."
        ),
        details=[
            "Employees are expected to act with integrity, honesty, and professionalism in all "
            "work-related activities and interactions.",
            "Conflicts of interest must be disclosed promptly to a supervisor or HR. Employees "
            "must not engage in activities that create or appear to create a conflict with the "
            "interests of the company.",
            "Confidential company information, client data, and proprietary materials must be "
            "protected at all times and must not be disclosed to unauthorized parties.",
            "Company resources — including equipment, systems, software, and working time — "
            "must be used responsibly and primarily for legitimate business purposes.",
            "Employees are expected to comply with all applicable laws and regulations, as well "
            "as all company policies and procedures.",
            "Workplace violence, threats, intimidation, or any behavior that endangers the "
            "safety of others will not be tolerated under any circumstances.",
        ],
        responsibilities=[
            "Conduct all business activities with honesty and in full compliance with applicable law.",
            "Disclose any actual or potential conflicts of interest promptly.",
            "Protect confidential information and company assets.",
            "Treat all colleagues, clients, and business partners with professionalism and respect.",
            "Report known or suspected violations of this policy to HR or management.",
        ],
        consequences=(
            "Violations of the Code of Conduct will result in disciplinary action commensurate "
            "with the severity of the offense, up to and including immediate termination. Certain "
            "violations may also be referred to law enforcement authorities."
        ),
    )


def policy_performance_review(company_name):
    return dict(
        title="Performance Review",
        purpose=(
            f"{company_name} is committed to the professional growth and development of every "
            f"employee. The performance review process provides a structured opportunity for "
            f"managers and employees to assess progress, set goals, and align on expectations."
        ),
        scope=(
            "This policy applies to all regular full-time and part-time employees. Contract "
            "workers and temporary staff are not covered under this policy unless otherwise "
            "specified in their agreement."
        ),
        details=[
            "Formal performance reviews are conducted on an annual basis, typically during Q4. "
            "New employees receive a 90-day check-in review within their first quarter of employment.",
            "Reviews are conducted by the employee's direct manager and may include input from "
            "peers, cross-functional partners, or senior leadership as appropriate.",
            "The review process evaluates performance against established goals, core competencies, "
            "and behavioral expectations. Outcomes may influence compensation decisions, promotion "
            "eligibility, and development planning.",
            "Managers are expected to provide ongoing, timely feedback throughout the year. "
            "Formal reviews should not be the only opportunity for performance dialogue.",
            "Employees who receive an unsatisfactory performance rating may be placed on a "
            "Performance Improvement Plan (PIP) outlining specific expectations and timelines "
            "for improvement.",
        ],
        responsibilities=[
            "Complete all required self-assessment materials prior to the review meeting.",
            "Engage honestly and constructively in the review discussion.",
            "Set and actively pursue professional development goals.",
            "Managers: deliver timely, specific, and actionable feedback throughout the year.",
        ],
        consequences=(
            "Failure to participate in the performance review process, or sustained performance "
            "below expectations without demonstrated improvement, may result in disciplinary "
            "action up to and including termination of employment."
        ),
    )


def policy_social_media(company_name):
    return dict(
        title="Social Media",
        purpose=(
            f"This policy provides guidance on the responsible use of social media by {company_name} "
            f"employees, both in a professional capacity and in personal use that may intersect "
            f"with the company's reputation or interests."
        ),
        scope=(
            "This policy applies to all employees and covers all social media platforms, including "
            "but not limited to LinkedIn, X (Twitter), Facebook, Instagram, TikTok, and any "
            "personal blogs or online forums."
        ),
        details=[
            f"Employees must not share confidential, proprietary, or sensitive information about "
            f"{company_name}, its clients, or its employees on any social media platform.",
            "When posting content that references the company or the employee's role, employees "
            "must make clear that opinions expressed are their own and do not represent the "
            "official position of the company.",
            "Employees are prohibited from making false, misleading, or defamatory statements "
            "about the company, its leadership, colleagues, competitors, or clients.",
            "Employees must not use social media to harass, bully, or discriminate against "
            "colleagues or any other individual.",
            f"Official social media accounts representing {company_name} may only be managed "
            f"by authorized personnel. Employees must not create unofficial accounts that could "
            f"be mistaken for official company channels.",
        ],
        responsibilities=[
            "Exercise professional judgment and discretion in all online activity that could be "
            "associated with the company.",
            "Immediately report to HR any social media content that may pose a legal or "
            "reputational risk to the company.",
            "Respect the privacy and confidentiality of colleagues when engaging online.",
        ],
        consequences=(
            "Violations of this policy — including unauthorized disclosure of confidential "
            "information or conduct that damages the company's reputation — may result in "
            "disciplinary action up to and including termination, and may expose the employee "
            "to personal legal liability."
        ),
    )


def policy_remote_work(company_name, state):
    return dict(
        title="Remote Work",
        purpose=(
            f"{company_name} recognizes that flexible work arrangements can enhance productivity "
            f"and employee well-being. This policy establishes expectations for employees who "
            f"work remotely on a full-time or part-time basis."
        ),
        scope=(
            "This policy applies to all employees who have been formally approved for remote "
            "work arrangements by their manager and Human Resources. Approval is required in "
            "advance and is not guaranteed for any position."
        ),
        details=[
            "Remote work arrangements must be approved in writing by the employee's manager "
            "and HR. Approval may be revoked at any time based on business needs or performance.",
            "Remote employees are expected to maintain the same level of productivity, "
            "availability, and professionalism as if working on-site. Core business hours "
            "must be observed unless an alternative schedule is approved.",
            "Employees are responsible for maintaining a safe, dedicated, and distraction-free "
            "workspace. The company is not responsible for personal utility or internet costs "
            "unless otherwise specified in a written agreement.",
            "Company-issued equipment must be used for all work activities. Employees must "
            "adhere to all IT security requirements, including VPN usage and endpoint protection.",
            f"Employees working remotely from outside {state} or outside the United States "
            f"must obtain prior written approval from HR, as this may have tax, benefits, and "
            f"legal implications for the company.",
        ],
        responsibilities=[
            "Be available and responsive during agreed-upon working hours.",
            "Maintain a secure, professional remote work environment.",
            "Protect company data and comply with all IT and data security policies.",
            "Communicate proactively with the manager regarding any issues affecting performance "
            "or availability.",
            "Attend required on-site meetings or events when requested by management.",
        ],
        consequences=(
            "Failure to comply with remote work expectations — including productivity issues, "
            "security breaches, or unauthorized changes to work location — may result in "
            "revocation of remote work privileges and further disciplinary action up to and "
            "including termination."
        ),
    )


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def build_acknowledgment_page(doc, company_name):
    add_page_break(doc)
    add_section_heading(doc, "Employee Acknowledgment")
    add_body(doc,
        f"By signing below, I acknowledge that I have received, read, and understand the "
        f"{company_name} HR Policy Handbook. I agree to comply with all policies and "
        f"understand that violations may result in disciplinary action, up to and including "
        f"termination of employment.",
        space_after=30
    )

    fields = [
        ("Employee Full Name (Print)", ""),
        ("Employee Signature", ""),
        ("Date", ""),
        ("Department", ""),
        ("Manager Name", ""),
    ]

    for label, _ in fields:
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        row = table.rows[0]
        lbl = row.cells[0].paragraphs[0]
        lbl.paragraph_format.space_before = Pt(4)
        lbl.paragraph_format.space_after = Pt(4)
        r = lbl.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = COLOR_TEXT

        val = row.cells[1].paragraphs[0]
        val.paragraph_format.space_before = Pt(4)
        val.paragraph_format.space_after = Pt(4)
        val.add_run(" ").font.size = Pt(10)
        doc.add_paragraph()

    add_paragraph(
        doc,
        f"Please return a signed copy of this page to Human Resources. A copy will be "
        f"retained in your personnel file.",
        font_size=9.5, italic=True, color=COLOR_MUTED, space_before=12
    )


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = COLOR_TEXT


def generate_policy_handbook(company_name, industry, company_size, state) -> Path:
    doc = Document()
    configure_document(doc)
    build_document_footer(doc)

    build_header(doc, company_name)
    build_intro(doc, company_name, industry, company_size, state)

    policies = [
        policy_eeo(company_name, state),
        policy_anti_harassment(company_name, state),
        policy_attendance(company_name, state),
        policy_pto(company_name, state),
        policy_code_of_conduct(company_name),
        policy_performance_review(company_name),
        policy_social_media(company_name),
        policy_remote_work(company_name, state),
    ]

    for i, policy in enumerate(policies, start=1):
        build_policy(
            doc, i,
            policy["title"],
            policy["purpose"],
            policy["scope"],
            policy["details"],
            policy["responsibilities"],
            policy["consequences"],
        )

    build_acknowledgment_page(doc, company_name)

    name_slug = company_name.replace(" ", "")
    date_slug = datetime.now().strftime("%Y-%m-%d")
    filename = f"{name_slug}_HR_Policies_{date_slug}.docx"

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = REPORTS_DIR / filename
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args():
    parser = argparse.ArgumentParser(
        description="Generate an HR Policy Handbook Word document.",
        epilog='Example: python policy_generator.py "Acme Corp" "Manufacturing" "250" "Ohio"'
    )
    parser.add_argument("company_name",  help="Name of the client company")
    parser.add_argument("industry",      help="Industry (e.g., Healthcare, Technology)")
    parser.add_argument("company_size",  help="Number of employees")
    parser.add_argument("state",         help="State of operation (for legal compliance)")
    return parser.parse_args()


def main():
    args = parse_args()

    print("\n" + "=" * 60)
    print(f"  {CONSULTANT_NAME}")
    print("  HR Policy Handbook Generator")
    print("=" * 60)
    print(f"  Company:  {args.company_name}")
    print(f"  Industry: {args.industry}")
    print(f"  Size:     {args.company_size} employees")
    print(f"  State:    {args.state}")
    print("=" * 60)
    print("  Generating handbook...\n")

    output_path = generate_policy_handbook(
        args.company_name,
        args.industry,
        args.company_size,
        args.state,
    )

    print("=" * 60)
    print("  HR Policy Handbook generated successfully.")
    print(f"  Saved to: {output_path}")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    main()
