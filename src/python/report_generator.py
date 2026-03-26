"""
report_generator.py
Waymark HR Group LLC

Generates a professional Executive Client Report as a Word document (.docx).
Saves the document to docs/reports/.

Usage:
    python report_generator.py "Company Name" "Industry" "engagement_type" ["Consultant Name"]

Arguments:
    1. Client company name        (required)
    2. Industry                   (required)
    3. Engagement type            (required) — onboarding | policy | workforce | full-audit
    4. Consultant name            (optional, default: Jamie Wahler)

Example:
    python report_generator.py "Acme Logistics" "Logistics" "full-audit" "Jamie Wahler"
"""

import sys
import argparse
from datetime import datetime, date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

COMPANY_NAME    = "Waymark HR Group LLC"
COMPANY_PHONE   = "716-225-6347"
COMPANY_EMAIL   = "Jamie.Wahler@waymarkhrgroup.com"
COMPANY_WEBSITE = "www.waymarkhrgroup.com"
COMPANY_ADDRESS = "Buffalo, NY"

COLOR_NAVY  = RGBColor(0x1A, 0x3A, 0x5C)
COLOR_TEAL  = RGBColor(0x00, 0x7A, 0x87)
COLOR_LGRAY = RGBColor(0xF4, 0xF6, 0xF8)
COLOR_TEXT  = RGBColor(0x2C, 0x2C, 0x2C)
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_MUTED = RGBColor(0x88, 0x88, 0x88)
COLOR_ACCENT = RGBColor(0xE8, 0xF4, 0xF6)

REPORTS_DIR = Path(__file__).resolve().parents[2] / "docs" / "reports"
LOGO_PATH   = Path(__file__).resolve().parents[2] / "docs" / "templates" / "Waymark_HR_Group_Logo_Full.png"

ENGAGEMENT_TITLES = {
    "onboarding":  "Employee Onboarding Engagement Report",
    "policy":      "HR Policy Development Engagement Report",
    "workforce":   "Workforce Analytics Engagement Report",
    "full-audit":  "Comprehensive HR Audit & Advisory Report",
}

VALID_ENGAGEMENTS = list(ENGAGEMENT_TITLES.keys())


# ---------------------------------------------------------------------------
# XML / style helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_bottom_border(paragraph, color="007A87", sz="6"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_top_border(paragraph, color="007A87", sz="4"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), sz)
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def add_paragraph(doc_or_cell, text="", bold=False, italic=False, font_size=11,
                  color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6):
    if hasattr(doc_or_cell, "add_paragraph"):
        p = doc_or_cell.add_paragraph()
    else:
        p = doc_or_cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        run.font.color.rgb = color or COLOR_TEXT
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEAL
    add_bottom_border(p)
    return p


def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_NAVY
    return p


def add_body(doc, text, italic=False, space_after=7):
    return add_paragraph(doc, text, font_size=10.5, italic=italic, space_after=space_after)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r1 = p.add_run(f"{bold_prefix} ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = COLOR_NAVY
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_TEXT
    return p


def add_page_break(doc):
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Page header (logo + rule on every page)
# ---------------------------------------------------------------------------

def build_page_header(doc):
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False

    # Clear default empty paragraph
    for p in header.paragraphs:
        p.clear()

    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)

    if LOGO_PATH.exists():
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), height=Inches(0.45))

    add_bottom_border(p, color="1A3A5C", sz="4")


# ---------------------------------------------------------------------------
# Page footer (contact info + page numbers)
# ---------------------------------------------------------------------------

def build_page_footer(doc, consultant_name):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    ft = footer.paragraphs[0]
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ft.paragraph_format.space_before = Pt(4)
    ft.paragraph_format.space_after = Pt(0)
    add_top_border(ft)

    run = ft.add_run(
        f"{COMPANY_NAME}  |  {consultant_name}  |  {COMPANY_PHONE}  |  {COMPANY_EMAIL}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED

    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)

    r_pre = p2.add_run("Page ")
    r_pre.font.size = Pt(8)
    r_pre.font.color.rgb = COLOR_MUTED

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    r_page = p2.add_run()
    r_page.font.size = Pt(8)
    r_page.font.color.rgb = COLOR_MUTED
    r_page._r.append(fldChar1)
    r_page._r.append(instrText)
    r_page._r.append(fldChar2)

    r_of = p2.add_run(" of ")
    r_of.font.size = Pt(8)
    r_of.font.color.rgb = COLOR_MUTED

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "begin")
    instrText2 = OxmlElement("w:instrText")
    instrText2.text = "NUMPAGES"
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    r_total = p2.add_run()
    r_total.font.size = Pt(8)
    r_total.font.color.rgb = COLOR_MUTED
    r_total._r.append(fldChar3)
    r_total._r.append(instrText2)
    r_total._r.append(fldChar4)


# ---------------------------------------------------------------------------
# Cover page
# ---------------------------------------------------------------------------

def build_cover_page(doc, client_name, engagement_type, consultant_name):
    engagement_title = ENGAGEMENT_TITLES[engagement_type]
    today = datetime.now().strftime("%B %d, %Y")

    # Logo centered
    if LOGO_PATH.exists():
        logo_p = doc.add_paragraph()
        logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_p.paragraph_format.space_before = Pt(40)
        logo_p.paragraph_format.space_after = Pt(30)
        logo_p.add_run().add_picture(str(LOGO_PATH), width=Inches(3.2))

    # Navy banner table
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1A3A5C")

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(20)
    p1.paragraph_format.space_after = Pt(6)
    r1 = p1.add_run(engagement_title.upper())
    r1.bold = True
    r1.font.size = Pt(15)
    r1.font.color.rgb = COLOR_WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run(f"Prepared for: {client_name}")
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0xB0, 0xC8, 0xD8)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(20)
    r3 = p3.add_run(f"Date Prepared: {today}  |  Prepared by: {consultant_name}")
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0xC8, 0xD8, 0xE8)

    # Confidentiality notice
    doc.add_paragraph()
    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conf.paragraph_format.space_before = Pt(30)
    r_conf = conf.add_run(
        "CONFIDENTIAL — This report has been prepared exclusively for the use of "
        f"{client_name} by {COMPANY_NAME}. It may not be reproduced or distributed "
        "without prior written consent."
    )
    r_conf.italic = True
    r_conf.font.size = Pt(9)
    r_conf.font.color.rgb = COLOR_MUTED

    add_page_break(doc)


# ---------------------------------------------------------------------------
# Executive summary
# ---------------------------------------------------------------------------

def build_executive_summary(doc, client_name, industry, engagement_type, consultant_name):
    add_section_heading(doc, "Executive Summary")

    engagement_overviews = {
        "onboarding": (
            f"Waymark HR Group LLC was engaged by {client_name} to design and deliver a "
            f"comprehensive employee onboarding program. The engagement focused on creating "
            f"a structured, consistent, and welcoming onboarding experience that reduces "
            f"time-to-productivity and improves new hire retention."
        ),
        "policy": (
            f"Waymark HR Group LLC was engaged by {client_name} to conduct a full review "
            f"of existing HR policies and develop a comprehensive HR Policy Handbook. The "
            f"engagement assessed current policy gaps, compliance risks, and alignment with "
            f"{industry} industry standards and applicable employment law."
        ),
        "workforce": (
            f"Waymark HR Group LLC was engaged by {client_name} to conduct an in-depth "
            f"workforce analytics review. The engagement analyzed headcount trends, employee "
            f"turnover, compensation benchmarking, and workforce planning to support data-driven "
            f"people strategy decisions."
        ),
        "full-audit": (
            f"Waymark HR Group LLC was engaged by {client_name} to perform a comprehensive "
            f"HR audit spanning onboarding, policy infrastructure, workforce analytics, and "
            f"overall HR operational effectiveness. This full-scope engagement provides "
            f"{client_name} with a complete picture of its current HR posture and a clear "
            f"roadmap for improvement."
        ),
    }

    findings = {
        "onboarding": [
            "New hire onboarding lacked a standardized structure, leading to inconsistent experiences across departments.",
            "First-week scheduling and IT setup processes were ad hoc and frequently delayed.",
            "No formal 30/60/90-day goal framework existed for new employees.",
            "Employee feedback indicated confusion around benefit enrollment timelines.",
        ],
        "policy": [
            "Multiple critical HR policies were absent or significantly outdated.",
            "Existing policies were not consistently communicated or acknowledged by employees.",
            f"Several practices were identified as potential compliance risks under {industry} regulations and applicable state law.",
            "No formal policy review cycle was in place to ensure ongoing accuracy.",
        ],
        "workforce": [
            "Voluntary turnover exceeded industry benchmarks in key departments.",
            "Compensation data revealed potential equity gaps in certain job families.",
            "Workforce planning was reactive rather than proactive, with limited headcount forecasting.",
            "Employee engagement indicators suggested risk of attrition among high performers.",
        ],
        "full-audit": [
            "Onboarding processes lacked standardization, resulting in inconsistent new hire experiences.",
            "Critical HR policies were absent, outdated, or not formally acknowledged by employees.",
            "Voluntary turnover and compensation equity gaps present retention and compliance risks.",
            "HR operations are largely manual, with limited use of technology to scale processes.",
            "No formal HR metrics dashboard exists to support leadership decision-making.",
        ],
    }

    impacts = {
        "onboarding": (
            f"A well-designed onboarding program is among the highest-ROI investments an "
            f"organization can make. Research consistently shows that structured onboarding "
            f"improves new hire retention by up to 82% and productivity by over 70%. The "
            f"program delivered to {client_name} establishes a strong foundation for employee "
            f"engagement from day one."
        ),
        "policy": (
            f"Clear, compliant HR policies reduce legal exposure, establish consistent expectations, "
            f"and support a fair and accountable workplace culture. The HR Policy Handbook "
            f"delivered to {client_name} provides a professionally structured framework aligned "
            f"with current employment law requirements and {industry} industry best practices."
        ),
        "workforce": (
            f"Data-driven workforce decisions reduce costs associated with turnover, misallocation "
            f"of talent, and reactive hiring. The analytics delivered through this engagement "
            f"equip {client_name}'s leadership with actionable intelligence to reduce attrition, "
            f"optimize compensation, and plan headcount with greater precision."
        ),
        "full-audit": (
            f"This comprehensive HR audit provides {client_name} with a clear, prioritized "
            f"roadmap to modernize its HR function. Addressing the findings identified in this "
            f"report will meaningfully reduce legal and operational risk, improve employee "
            f"experience, and position the organization for scalable, sustainable growth."
        ),
    }

    add_subsection_heading(doc, "Overview")
    add_body(doc, engagement_overviews[engagement_type])

    add_subsection_heading(doc, "Key Findings")
    for finding in findings[engagement_type]:
        add_bullet(doc, finding)

    add_subsection_heading(doc, "Business Impact")
    add_body(doc, impacts[engagement_type])

    add_page_break(doc)


# ---------------------------------------------------------------------------
# Services delivered
# ---------------------------------------------------------------------------

def build_services_delivered(doc, client_name, engagement_type):
    add_section_heading(doc, "Services Delivered")

    deliverables = {
        "onboarding": [
            ("Onboarding Program Design",       "Developed a structured onboarding framework including pre-boarding, Day 1, and first-week workflows."),
            ("Welcome Package Creation",         "Produced a branded, personalized Employee Onboarding Welcome Package (.docx) for each new hire."),
            ("First-Week Schedule",              "Designed a detailed Day 1–5 schedule covering orientation, IT setup, team introductions, and goal-setting."),
            ("IT Setup Checklist",               "Created a comprehensive IT onboarding checklist covering hardware, accounts, security, and collaboration tools."),
            ("HR Policy Summary",                "Delivered a concise summary of key HR policies for inclusion in new hire materials."),
            ("Emergency Contacts Directory",     "Compiled an internal contacts and emergency services reference page for all onboarding packages."),
        ],
        "policy": [
            ("Policy Gap Assessment",            "Reviewed existing documentation to identify missing, outdated, or non-compliant HR policies."),
            ("HR Policy Handbook",               "Drafted a comprehensive, state-specific HR Policy Handbook covering 8 core policy areas."),
            ("EEO & Anti-Harassment Policies",   "Developed compliant Equal Employment Opportunity and Anti-Harassment policies aligned with federal and state law."),
            ("Attendance & PTO Policies",        "Drafted Attendance, Punctuality, and Paid Time Off policies tailored to the organization's needs."),
            ("Code of Conduct",                  "Created a professional Code of Conduct establishing behavioral and ethical standards."),
            ("Compliance Review",                "Assessed policy content against applicable employment law requirements."),
        ],
        "workforce": [
            ("Headcount Analysis",               "Analyzed current workforce composition, growth trends, and departmental distribution."),
            ("Turnover & Retention Analysis",    "Evaluated voluntary and involuntary turnover rates against industry benchmarks."),
            ("Compensation Benchmarking",        "Compared internal compensation structures to market data across key job families."),
            ("Workforce Planning Review",        "Assessed current headcount planning processes and identified forecasting gaps."),
            ("Engagement Risk Assessment",       "Identified indicators of disengagement and flight risk among key employee segments."),
            ("HR Metrics Dashboard Design",      "Defined core KPIs and recommended a metrics framework for ongoing HR reporting."),
        ],
        "full-audit": [
            ("HR Operations Assessment",         "Conducted a full review of HR processes, systems, and documentation across all functional areas."),
            ("Onboarding Program Redesign",      "Delivered a structured onboarding program with branded new hire welcome packages."),
            ("HR Policy Handbook",               "Developed a comprehensive, compliant HR Policy Handbook covering 8 core policy areas."),
            ("Workforce Analytics Review",       "Performed headcount, turnover, compensation, and engagement analyses."),
            ("Compliance Risk Assessment",       "Identified policy and practice gaps posing potential legal or regulatory exposure."),
            ("HR Roadmap Development",           "Produced a prioritized 90-day and long-term HR improvement roadmap."),
        ],
    }

    tools = {
        "onboarding":  ["Waymark Onboarding Generator (Python)", "Microsoft Word (.docx)", "HR policy framework templates"],
        "policy":      ["Waymark Policy Generator (Python)", "Federal and state employment law reference", "SHRM policy framework"],
        "workforce":   ["Workforce data analysis toolkit", "Compensation benchmarking databases", "HR metrics framework"],
        "full-audit":  ["Waymark Onboarding Generator (Python)", "Waymark Policy Generator (Python)", "Workforce analytics toolkit", "HR audit framework"],
    }

    timeline = {
        "onboarding": [
            ("Week 1", "Discovery call — roles, departments, hiring cadence"),
            ("Week 2", "Onboarding framework design and review"),
            ("Week 3", "Welcome package development and delivery"),
            ("Week 4", "Final review, handoff, and implementation guidance"),
        ],
        "policy": [
            ("Week 1", "Policy gap assessment and discovery"),
            ("Week 2", "Policy drafting — EEO, harassment, attendance, PTO"),
            ("Week 3", "Policy drafting — code of conduct, performance, social media, remote work"),
            ("Week 4", "Review cycle, compliance check, and final handbook delivery"),
        ],
        "workforce": [
            ("Week 1", "Data collection and workforce audit"),
            ("Week 2", "Turnover and retention analysis"),
            ("Week 3", "Compensation benchmarking and engagement assessment"),
            ("Week 4", "Findings synthesis, metrics framework, and report delivery"),
        ],
        "full-audit": [
            ("Week 1",   "Full HR audit — discovery, data collection, document review"),
            ("Week 2",   "Onboarding program design and policy drafting"),
            ("Week 3",   "Workforce analytics, compensation review, and compliance assessment"),
            ("Week 4",   "Roadmap development, report preparation"),
            ("Week 5",   "Final review, delivery, and leadership presentation"),
        ],
    }

    add_subsection_heading(doc, "Deliverables Completed")
    for item, description in deliverables[engagement_type]:
        add_bullet(doc, description, bold_prefix=f"{item}:")

    add_subsection_heading(doc, "Engagement Timeline")
    tl = doc.add_table(rows=1 + len(timeline[engagement_type]), cols=2)
    tl.style = "Table Grid"
    tl.alignment = WD_TABLE_ALIGNMENT.LEFT

    for col_idx, header in enumerate(["Phase", "Activities"]):
        cell = tl.cell(0, col_idx)
        set_cell_bg(cell, "1A3A5C")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE

    for row_idx, (phase, activity) in enumerate(timeline[engagement_type], start=1):
        row = tl.rows[row_idx]
        bg = "F4F6F8" if row_idx % 2 == 0 else "FFFFFF"

        phase_cell = row.cells[0]
        set_cell_bg(phase_cell, "E8F4F6")
        pp = phase_cell.paragraphs[0]
        pp.paragraph_format.space_before = Pt(3)
        pp.paragraph_format.space_after = Pt(3)
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr = pp.add_run(phase)
        pr.bold = True
        pr.font.size = Pt(10)
        pr.font.color.rgb = COLOR_NAVY

        act_cell = row.cells[1]
        set_cell_bg(act_cell, bg)
        ap = act_cell.paragraphs[0]
        ap.paragraph_format.space_before = Pt(3)
        ap.paragraph_format.space_after = Pt(3)
        ar = ap.add_run(activity)
        ar.font.size = Pt(10)
        ar.font.color.rgb = COLOR_TEXT

    doc.add_paragraph()
    add_subsection_heading(doc, "Tools & Frameworks Used")
    for tool in tools[engagement_type]:
        add_bullet(doc, tool)

    add_page_break(doc)


# ---------------------------------------------------------------------------
# Key recommendations
# ---------------------------------------------------------------------------

def build_recommendations(doc, client_name, engagement_type):
    add_section_heading(doc, "Key Recommendations")
    add_body(doc,
        f"The following five recommendations are prioritized based on potential business impact "
        f"and implementation effort. Each is rated on a High / Medium / Low scale.",
        italic=True
    )

    recs = {
        "onboarding": [
            ("Implement a Formal Pre-Boarding Process",
             "Engage new hires before Day 1 with welcome emails, paperwork, and logistical guidance to reduce first-day friction.",
             "High", "Low", "Quick Win"),
            ("Assign Onboarding Buddies",
             "Pair each new hire with a peer mentor for the first 60 days to accelerate cultural integration and reduce isolation.",
             "High", "Low", "Quick Win"),
            ("Establish a 30/60/90-Day Goal Framework",
             "Define clear performance and learning milestones for all new hires, reviewed jointly with their manager.",
             "High", "Medium", "Short-Term"),
            ("Automate New Hire Paperwork",
             "Implement a digital HR platform or e-signature solution to eliminate manual onboarding paperwork.",
             "Medium", "Medium", "Short-Term"),
            ("Conduct Structured 30-Day New Hire Surveys",
             "Collect feedback at the 30-day mark to identify onboarding gaps and continuously improve the program.",
             "Medium", "Low", "Quick Win"),
        ],
        "policy": [
            ("Distribute and Obtain Signed Policy Acknowledgments",
             "Ensure all employees receive, read, and sign acknowledgment of the new HR Policy Handbook immediately.",
             "High", "Low", "Quick Win"),
            ("Schedule Annual Policy Reviews",
             "Establish a formal calendar-based policy review process to keep documentation current with evolving law.",
             "High", "Low", "Quick Win"),
            ("Conduct Manager-Level Policy Training",
             "Train all people managers on key policies — especially harassment, attendance, and performance — to ensure consistent enforcement.",
             "High", "Medium", "Short-Term"),
            ("Implement a Digital Policy Distribution System",
             "Move policy acknowledgment and updates to a digital HR platform to ensure version control and audit trails.",
             "Medium", "Medium", "Short-Term"),
            ("Develop a Whistleblower & Reporting Mechanism",
             "Establish a confidential reporting channel for policy violations and workplace concerns.",
             "High", "Medium", "Long-Term Investment"),
        ],
        "workforce": [
            ("Conduct Stay Interviews with High Performers",
             "Proactively identify retention risks by asking top performers what would cause them to leave — and address it.",
             "High", "Low", "Quick Win"),
            ("Address Compensation Equity Gaps",
             "Review and remediate identified pay disparities within affected job families to reduce legal exposure and improve morale.",
             "High", "Medium", "Short-Term"),
            ("Implement Quarterly Headcount Planning Reviews",
             "Shift from reactive to proactive hiring by building a quarterly workforce planning cadence with department heads.",
             "High", "Medium", "Short-Term"),
            ("Launch an HR Metrics Dashboard",
             "Implement a real-time HR dashboard tracking turnover, headcount, time-to-fill, and engagement KPIs.",
             "Medium", "High", "Long-Term Investment"),
            ("Develop a Succession Planning Framework",
             "Identify critical roles and high-potential employees, and create development plans to build internal pipeline.",
             "High", "High", "Long-Term Investment"),
        ],
        "full-audit": [
            ("Standardize and Launch the New Onboarding Program",
             "Roll out the newly designed onboarding program company-wide immediately to improve new hire retention from day one.",
             "High", "Low", "Quick Win"),
            ("Distribute the HR Policy Handbook and Collect Acknowledgments",
             "Issue the completed HR Policy Handbook to all employees and obtain signed acknowledgments within 30 days.",
             "High", "Low", "Quick Win"),
            ("Remediate Compensation Equity Gaps",
             "Address identified pay disparities in key job families to reduce legal risk and improve employee trust.",
             "High", "Medium", "Short-Term"),
            ("Invest in an HRIS Platform",
             "Implement a Human Resource Information System to centralize employee data, automate workflows, and enable reporting.",
             "High", "High", "Long-Term Investment"),
            ("Build an HR Metrics & Reporting Function",
             "Establish a regular HR reporting cadence with leadership using the defined KPI framework from this engagement.",
             "Medium", "Medium", "Short-Term"),
        ],
    }

    impact_colors = {"High": "D6F0D6", "Medium": "FFF3CD", "Low": "FFE0CC"}
    type_colors   = {"Quick Win": "D6F0D6", "Short-Term": "FFF3CD", "Long-Term Investment": "FFE0CC"}

    rec_list = recs[engagement_type]
    table = doc.add_table(rows=1 + len(rec_list), cols=5)
    table.style = "Table Grid"

    headers = ["#", "Recommendation", "Impact", "Effort", "Type"]
    widths  = [Inches(0.3), Inches(3.2), Inches(0.85), Inches(0.85), Inches(1.3)]
    for col_idx, (header, width) in enumerate(zip(headers, widths)):
        cell = table.cell(0, col_idx)
        set_cell_bg(cell, "1A3A5C")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_WHITE
        table.columns[col_idx].width = width

    for row_idx, (title, desc, impact, effort, rtype) in enumerate(rec_list, start=1):
        row = table.rows[row_idx]
        bg = "F4F6F8" if row_idx % 2 == 0 else "FFFFFF"

        # Number
        num_cell = row.cells[0]
        set_cell_bg(num_cell, bg)
        np = num_cell.paragraphs[0]
        np.alignment = WD_ALIGN_PARAGRAPH.CENTER
        np.paragraph_format.space_before = Pt(4)
        np.paragraph_format.space_after = Pt(4)
        nr = np.add_run(str(row_idx))
        nr.bold = True
        nr.font.size = Pt(10)
        nr.font.color.rgb = COLOR_NAVY

        # Title + description
        desc_cell = row.cells[1]
        set_cell_bg(desc_cell, bg)
        dp = desc_cell.paragraphs[0]
        dp.paragraph_format.space_before = Pt(4)
        dp.paragraph_format.space_after = Pt(2)
        dr = dp.add_run(f"{title}\n")
        dr.bold = True
        dr.font.size = Pt(10)
        dr.font.color.rgb = COLOR_NAVY
        dr2 = dp.add_run(desc)
        dr2.font.size = Pt(9.5)
        dr2.font.color.rgb = COLOR_TEXT

        # Impact
        imp_cell = row.cells[2]
        set_cell_bg(imp_cell, impact_colors.get(impact, bg))
        ip = imp_cell.paragraphs[0]
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.paragraph_format.space_before = Pt(4)
        ip.paragraph_format.space_after = Pt(4)
        ir = ip.add_run(impact)
        ir.bold = True
        ir.font.size = Pt(10)
        ir.font.color.rgb = COLOR_TEXT

        # Effort
        eff_cell = row.cells[3]
        set_cell_bg(eff_cell, impact_colors.get(effort, bg))
        ep = eff_cell.paragraphs[0]
        ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ep.paragraph_format.space_before = Pt(4)
        ep.paragraph_format.space_after = Pt(4)
        er = ep.add_run(effort)
        er.bold = True
        er.font.size = Pt(10)
        er.font.color.rgb = COLOR_TEXT

        # Type
        type_cell = row.cells[4]
        set_cell_bg(type_cell, type_colors.get(rtype, bg))
        tp = type_cell.paragraphs[0]
        tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tp.paragraph_format.space_before = Pt(4)
        tp.paragraph_format.space_after = Pt(4)
        tr = tp.add_run(rtype)
        tr.bold = True
        tr.font.size = Pt(9.5)
        tr.font.color.rgb = COLOR_TEXT

    doc.add_paragraph()
    add_page_break(doc)


# ---------------------------------------------------------------------------
# Next steps / 30-60-90 plan
# ---------------------------------------------------------------------------

def build_next_steps(doc, client_name, engagement_type, consultant_name):
    add_section_heading(doc, "Next Steps & 30/60/90-Day Action Plan")

    followon = {
        "onboarding": [
            "HR Policy Handbook development to complement the onboarding program",
            "New hire survey design and 30-day check-in process implementation",
            "Manager onboarding training and facilitation",
            "Annual onboarding program review and refresh",
        ],
        "policy": [
            "Manager training on policy enforcement and documentation",
            "Employee handbook digital distribution and acknowledgment tracking",
            "Workforce analytics engagement to identify retention and compensation risks",
            "Annual compliance review to keep policies current with legal changes",
        ],
        "workforce": [
            "Compensation equity remediation planning and implementation support",
            "Succession planning framework development",
            "HR metrics dashboard implementation",
            "Engagement survey design and administration",
        ],
        "full-audit": [
            "HRIS platform evaluation and implementation support",
            "Manager leadership development and HR training program",
            "Ongoing HR advisory retainer for continued compliance and strategy support",
            "Annual HR audit to track progress against this report's recommendations",
        ],
    }

    plan = {
        "onboarding": [
            ("30 Days",  [
                "Distribute onboarding materials to all hiring managers",
                "Train managers on the new onboarding program and expectations",
                "Launch pre-boarding process for all incoming hires",
                "Assign onboarding buddies for current open positions",
            ]),
            ("60 Days",  [
                "Conduct 30-day new hire surveys and review results",
                "Refine onboarding schedule based on initial feedback",
                "Begin 30/60/90-day goal-setting conversations with recent hires",
                "Evaluate IT setup checklist completion rates",
            ]),
            ("90 Days",  [
                "Complete first full onboarding cycle review",
                "Present onboarding effectiveness report to leadership",
                "Identify opportunities for automation and digital tools",
                "Plan for ongoing program maintenance and annual refresh",
            ]),
        ],
        "policy": [
            ("30 Days",  [
                "Distribute HR Policy Handbook to all employees",
                "Collect signed acknowledgments from all staff",
                "Brief managers on key policy updates and enforcement expectations",
                "Post policies on company intranet or shared drive",
            ]),
            ("60 Days",  [
                "Conduct manager policy training sessions",
                "Confirm 100% acknowledgment receipt from all employees",
                "Address any employee questions or policy clarifications",
                "Begin evaluation of digital policy management platforms",
            ]),
            ("90 Days",  [
                "Establish annual policy review calendar",
                "Assess policy adherence and any early compliance issues",
                "Develop a whistleblower and reporting mechanism",
                "Plan next HR engagement (workforce analytics recommended)",
            ]),
        ],
        "workforce": [
            ("30 Days",  [
                "Share workforce analytics findings with senior leadership",
                "Initiate compensation equity review with HR and Finance",
                "Launch stay interview program with high-performer population",
                "Define KPIs for ongoing HR metrics tracking",
            ]),
            ("60 Days",  [
                "Present compensation remediation recommendations to leadership",
                "Implement quarterly headcount planning cadence",
                "Begin HR metrics dashboard design and configuration",
                "Conduct first round of stay interviews and synthesize themes",
            ]),
            ("90 Days",  [
                "Finalize and implement compensation adjustments",
                "Launch HR metrics dashboard for leadership review",
                "Develop succession planning framework for critical roles",
                "Present 90-day workforce progress report to leadership",
            ]),
        ],
        "full-audit": [
            ("30 Days",  [
                "Launch new onboarding program for all incoming hires",
                "Distribute HR Policy Handbook and collect employee acknowledgments",
                "Share audit findings with senior leadership",
                "Initiate compensation equity review with HR and Finance",
            ]),
            ("60 Days",  [
                "Complete manager training on policies and onboarding program",
                "Begin HRIS platform evaluation process",
                "Implement quarterly headcount planning cadence",
                "Launch stay interview program with high-performer population",
            ]),
            ("90 Days",  [
                "Finalize compensation remediation plan and begin implementation",
                "Select and initiate HRIS implementation project",
                "Launch HR metrics dashboard for leadership",
                "Present 90-day progress report against audit recommendations",
            ]),
        ],
    }

    add_subsection_heading(doc, "Proposed Follow-On Services")
    add_body(doc,
        f"Based on the findings and outcomes of this engagement, Waymark HR Group LLC "
        f"recommends the following follow-on services to continue building {client_name}'s "
        f"HR capability:"
    )
    for item in followon[engagement_type]:
        add_bullet(doc, item)

    add_subsection_heading(doc, "30/60/90-Day Action Plan")
    for phase, actions in plan[engagement_type]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(phase)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = COLOR_TEAL
        for action in actions:
            add_bullet(doc, action)

    add_page_break(doc)


# ---------------------------------------------------------------------------
# Closing page
# ---------------------------------------------------------------------------

def build_closing_page(doc, client_name, consultant_name):
    add_section_heading(doc, "Closing")

    add_paragraph(doc, f"Dear {client_name} Team,", bold=True, font_size=10.5, space_after=10)

    closing_paragraphs = [
        f"On behalf of everyone at {COMPANY_NAME}, thank you for the opportunity to support "
        f"your organization through this engagement. It has been a privilege to work alongside "
        f"your team, and we are proud of the work we have accomplished together.",

        "The deliverables contained in this report represent a meaningful investment in your "
        "people and your organization's future. We are confident that implementing the "
        "recommendations outlined here will strengthen your HR function, reduce risk, and "
        "create a better experience for your employees.",

        "We remain committed to your success beyond this engagement. Whether you need support "
        "implementing these recommendations, have questions about the deliverables, or are "
        "ready to begin the next phase of your HR journey, our team is here to help.",

        "Please do not hesitate to reach out at any time. We look forward to continuing to "
        "be a trusted partner for your organization.",
    ]

    for para in closing_paragraphs:
        add_body(doc, para, space_after=10)

    add_body(doc, "Warm regards,", space_after=2)
    add_paragraph(doc, consultant_name, bold=True, font_size=10.5, space_after=2)
    add_body(doc, COMPANY_NAME, space_after=2)

    # Contact info box
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_bg(cell, "F4F6F8")

    lines = [
        ("Phone:",   COMPANY_PHONE),
        ("Email:",   COMPANY_EMAIL),
        ("Website:", COMPANY_WEBSITE),
        ("Address:", COMPANY_ADDRESS),
    ]

    for i, (label, value) in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(6) if i == 0 else Pt(3)
        p.paragraph_format.space_after = Pt(3) if i < len(lines) - 1 else Pt(10)
        r1 = p.add_run(f"{label}  ")
        r1.bold = True
        r1.font.size = Pt(10.5)
        r1.font.color.rgb = COLOR_NAVY
        r2 = p.add_run(value)
        r2.font.size = Pt(10.5)
        r2.font.color.rgb = COLOR_TEAL

    doc.add_paragraph()
    add_body(doc,
        "This report is confidential and intended solely for the use of the named client. "
        f"© {datetime.now().year} {COMPANY_NAME}. All rights reserved.",
        italic=True, space_after=4
    )


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def configure_document(doc):
    section = doc.sections[0]
    section.top_margin    = Inches(1.1)
    section.bottom_margin = Inches(0.9)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = COLOR_TEXT


# ---------------------------------------------------------------------------
# Main document assembly
# ---------------------------------------------------------------------------

def generate_report(client_name, industry, engagement_type, consultant_name) -> Path:
    doc = Document()
    configure_document(doc)
    build_page_header(doc)
    build_page_footer(doc, consultant_name)

    build_cover_page(doc, client_name, engagement_type, consultant_name)
    build_executive_summary(doc, client_name, industry, engagement_type, consultant_name)
    build_services_delivered(doc, client_name, engagement_type)
    build_recommendations(doc, client_name, engagement_type)
    build_next_steps(doc, client_name, engagement_type, consultant_name)
    build_closing_page(doc, client_name, consultant_name)

    name_slug = client_name.replace(" ", "")
    date_slug = datetime.now().strftime("%Y-%m-%d")
    filename = f"{name_slug}_ClientReport_{date_slug}.docx"

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = REPORTS_DIR / filename
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args():
    parser = argparse.ArgumentParser(
        description="Generate an Executive Client Report for a Waymark HR Group engagement.",
        epilog='Example: python report_generator.py "Acme Corp" "Manufacturing" "full-audit" "Jamie Wahler"'
    )
    parser.add_argument("company_name",    help="Client company name")
    parser.add_argument("industry",        help="Client industry (e.g., Healthcare, Logistics)")
    parser.add_argument("engagement_type", choices=VALID_ENGAGEMENTS,
                        help="Type of engagement: onboarding | policy | workforce | full-audit")
    parser.add_argument("consultant_name", nargs="?", default="Jamie Wahler",
                        help="Consultant name (default: Jamie Wahler)")
    return parser.parse_args()


def main():
    args = parse_args()

    print("\n" + "=" * 60)
    print(f"  {COMPANY_NAME}")
    print("  Executive Client Report Generator")
    print("=" * 60)
    print(f"  Client:     {args.company_name}")
    print(f"  Industry:   {args.industry}")
    print(f"  Engagement: {args.engagement_type}")
    print(f"  Consultant: {args.consultant_name}")
    print("=" * 60)
    print("  Generating report...\n")

    output_path = generate_report(
        args.company_name,
        args.industry,
        args.engagement_type,
        args.consultant_name,
    )

    print("=" * 60)
    print("  Executive Client Report generated successfully.")
    print(f"  Saved to: {output_path}")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    main()
