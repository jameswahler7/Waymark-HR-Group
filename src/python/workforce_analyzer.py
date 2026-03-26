"""
workforce_analyzer.py
Waymark HR Group LLC

Analyzes workforce data from a CSV or Excel file and generates a professional
Word document report covering headcount, compensation, turnover risk, and
HR recommendations.

Usage:
    python workforce_analyzer.py --input PATH_TO_FILE --company "Company Name"

Arguments:
    --input       Path to a CSV or Excel (.xlsx / .xls) workforce data file
    --company     Client company name (used in the report title and filename)

Expected columns (case-insensitive, spaces or underscores accepted):
    employee_id, first_name, last_name, department, job_title,
    employment_type, hire_date, annual_salary, last_performance_review

Output:
    docs/reports/<CompanyName>_Workforce_Analysis_YYYY-MM-DD.docx
"""

import sys
import argparse
from datetime import date, datetime
from pathlib import Path

try:
    import pandas as pd
except ImportError:
    print("Error: pandas is required. Install it with: pip install pandas")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

CONSULTANT_NAME    = "Waymark HR Group LLC"
CONSULTANT_PHONE   = "716-225-6347"
CONSULTANT_EMAIL   = "Jamie.Wahler@waymarkhrgroup.com"
CONSULTANT_WEBSITE = "www.waymarkhrgroup.com"

COLOR_NAVY       = RGBColor(0x1A, 0x3A, 0x5C)
COLOR_TEAL       = RGBColor(0x00, 0x7A, 0x87)
COLOR_LIGHT_GRAY = RGBColor(0xF4, 0xF6, 0xF8)
COLOR_TEXT       = RGBColor(0x2C, 0x2C, 0x2C)
COLOR_WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_MUTED      = RGBColor(0x88, 0x88, 0x88)
COLOR_FLAG       = RGBColor(0xC0, 0x39, 0x2B)   # red — used for flagged items

REPORTS_DIR = Path(__file__).resolve().parents[2] / "docs" / "reports"
LOGO_PATH   = Path(__file__).resolve().parents[2] / "docs" / "templates" / "Waymark_HR_Group_Logo_Full.png"

# Pay equity: flag a role when the spread between highest and lowest salary
# within the same job title exceeds this percentage.
PAY_EQUITY_THRESHOLD = 0.15   # 15 %

# Turnover risk: flag employees hired within this many months of today.
NEW_HIRE_MONTHS = 6

# Turnover concentration: flag a department when this fraction or more of its
# headcount consists of new hires.
HIGH_NEW_HIRE_RATIO = 0.30    # 30 %


# ---------------------------------------------------------------------------
# Column normalisation
# ---------------------------------------------------------------------------

COLUMN_ALIASES = {
    "employee_id":              ["employee_id", "employeeid", "id", "emp_id", "empid"],
    "first_name":               ["first_name", "firstname", "first"],
    "last_name":                ["last_name", "lastname", "last"],
    "department":               ["department", "dept"],
    "job_title":                ["job_title", "jobtitle", "title", "position", "role"],
    "employment_type":          ["employment_type", "employmenttype", "type", "status",
                                 "employment status", "emp_type"],
    "hire_date":                ["hire_date", "hiredate", "start_date", "startdate",
                                 "date_hired", "datehired"],
    "annual_salary":            ["annual_salary", "annualsalary", "salary", "base_salary",
                                 "basesalary", "compensation", "pay"],
    "last_performance_review":  ["last_performance_review", "lastperformancereview",
                                 "performance_review", "review_date", "reviewdate",
                                 "last_review"],
}


def normalise_columns(df: pd.DataFrame) -> pd.DataFrame:
    """Rename DataFrame columns to canonical snake_case names."""
    rename_map = {}
    lowered = {col.lower().replace(" ", "_"): col for col in df.columns}
    for canonical, aliases in COLUMN_ALIASES.items():
        for alias in aliases:
            if alias in lowered:
                rename_map[lowered[alias]] = canonical
                break
    return df.rename(columns=rename_map)


def validate_columns(df: pd.DataFrame) -> list[str]:
    """Return a list of required columns that are missing."""
    required = ["department", "employment_type", "hire_date", "annual_salary"]
    return [col for col in required if col not in df.columns]


# ---------------------------------------------------------------------------
# Data loading
# ---------------------------------------------------------------------------

def load_data(file_path: str) -> pd.DataFrame:
    path = Path(file_path)
    if not path.exists():
        print(f"Error: File not found — {file_path}")
        sys.exit(1)

    suffix = path.suffix.lower()
    if suffix == ".csv":
        df = pd.read_csv(path)
    elif suffix in (".xlsx", ".xls"):
        df = pd.read_excel(path)
    else:
        print(f"Error: Unsupported file type '{suffix}'. Use .csv, .xlsx, or .xls.")
        sys.exit(1)

    df = normalise_columns(df)

    missing = validate_columns(df)
    if missing:
        print(f"Error: Required column(s) not found in data: {', '.join(missing)}")
        print("Please ensure your file contains the following columns:")
        print("  department, employment_type, hire_date, annual_salary")
        sys.exit(1)

    # Parse dates
    df["hire_date"] = pd.to_datetime(df["hire_date"], errors="coerce")
    if "last_performance_review" in df.columns:
        df["last_performance_review"] = pd.to_datetime(
            df["last_performance_review"], errors="coerce"
        )

    # Compute tenure in years (as float) and months (as int)
    today = pd.Timestamp(date.today())
    df["tenure_years"] = (today - df["hire_date"]).dt.days / 365.25
    df["tenure_months"] = ((today - df["hire_date"]).dt.days / 30.44).astype(int)

    # Normalise employment type to Full-Time / Part-Time
    def normalise_type(val):
        if pd.isna(val):
            return "Unknown"
        v = str(val).lower().strip()
        if "full" in v or v in ("ft", "f"):
            return "Full-Time"
        if "part" in v or v in ("pt", "p"):
            return "Part-Time"
        return str(val).strip()

    df["employment_type"] = df["employment_type"].apply(normalise_type)

    # Build a display name if first/last name columns exist
    if "first_name" in df.columns and "last_name" in df.columns:
        df["full_name"] = df["first_name"].fillna("") + " " + df["last_name"].fillna("")
        df["full_name"] = df["full_name"].str.strip()
    elif "full_name" not in df.columns:
        df["full_name"] = df.get("employee_id", pd.Series(["—"] * len(df)))

    return df


# ---------------------------------------------------------------------------
# Analysis
# ---------------------------------------------------------------------------

def analyze_headcount(df: pd.DataFrame) -> dict:
    """Return headcount summary metrics."""
    total_by_dept = (
        df.groupby("department", sort=True)
        .size()
        .reset_index(name="total")
    )

    ft_pt_by_dept = (
        df.groupby(["department", "employment_type"], sort=True)
        .size()
        .unstack(fill_value=0)
        .reset_index()
    )

    # Ensure both columns exist even if data has only one type
    for col in ("Full-Time", "Part-Time"):
        if col not in ft_pt_by_dept.columns:
            ft_pt_by_dept[col] = 0

    avg_tenure_by_dept = (
        df.groupby("department", sort=True)["tenure_years"]
        .mean()
        .reset_index()
    )
    avg_tenure_by_dept.columns = ["department", "avg_tenure_years"]

    return {
        "total_employees":   len(df),
        "total_by_dept":     total_by_dept,
        "ft_pt_by_dept":     ft_pt_by_dept,
        "avg_tenure_by_dept": avg_tenure_by_dept,
    }


def analyze_compensation(df: pd.DataFrame) -> dict:
    """Return compensation metrics and pay equity flags."""
    avg_by_dept = (
        df.groupby("department", sort=True)["annual_salary"]
        .mean()
        .reset_index()
    )
    avg_by_dept.columns = ["department", "avg_salary"]

    # Salary range per job title (only titles with at least 1 employee)
    salary_by_role = (
        df.groupby("job_title", sort=True)["annual_salary"]
        .agg(count="count", min_salary="min", max_salary="max", avg_salary="mean")
        .reset_index()
    )

    # Pay equity flags: roles with >= 2 employees where spread > threshold
    equity_flags = []
    for _, row in salary_by_role.iterrows():
        if row["count"] >= 2 and row["min_salary"] > 0:
            spread = (row["max_salary"] - row["min_salary"]) / row["min_salary"]
            if spread > PAY_EQUITY_THRESHOLD:
                equity_flags.append({
                    "job_title":  row["job_title"],
                    "count":      int(row["count"]),
                    "min_salary": row["min_salary"],
                    "max_salary": row["max_salary"],
                    "spread_pct": spread * 100,
                })

    return {
        "avg_by_dept":    avg_by_dept,
        "salary_by_role": salary_by_role,
        "equity_flags":   equity_flags,
    }


def analyze_turnover_risk(df: pd.DataFrame) -> dict:
    """Return turnover risk indicators."""
    new_hires = df[df["tenure_months"] < NEW_HIRE_MONTHS].copy()
    new_hires_sorted = new_hires.sort_values("tenure_months")

    # Departments where >= HIGH_NEW_HIRE_RATIO of headcount are new hires
    dept_totals    = df.groupby("department").size()
    dept_new_hires = new_hires.groupby("department").size()
    dept_ratio     = (dept_new_hires / dept_totals).dropna()
    high_risk_depts = dept_ratio[dept_ratio >= HIGH_NEW_HIRE_RATIO].reset_index()
    high_risk_depts.columns = ["department", "new_hire_ratio"]
    high_risk_depts["new_hire_count"] = high_risk_depts["department"].map(
        dept_new_hires
    ).astype(int)
    high_risk_depts["total"] = high_risk_depts["department"].map(dept_totals).astype(int)

    # Missing performance review dates
    missing_review = pd.DataFrame()
    if "last_performance_review" in df.columns:
        missing_review = df[df["last_performance_review"].isna()].copy()

    return {
        "new_hires":         new_hires_sorted,
        "high_risk_depts":   high_risk_depts,
        "missing_review":    missing_review,
    }


def generate_recommendations(
    headcount: dict,
    compensation: dict,
    turnover: dict,
) -> list[str]:
    """Derive 3–5 action-item bullet points from analysis results."""
    recs = []

    # Recommendation 1 — new hires / onboarding
    new_hire_count = len(turnover["new_hires"])
    if new_hire_count > 0:
        recs.append(
            f"Implement a structured 30-60-90 day onboarding program to support the "
            f"{new_hire_count} employee(s) hired within the past six months. Early engagement "
            f"programs significantly reduce first-year voluntary turnover."
        )

    # Recommendation 2 — pay equity
    if compensation["equity_flags"]:
        roles_flagged = ", ".join(
            f['job_title'] for f in compensation["equity_flags"]
        )
        recs.append(
            f"Conduct a formal compensation equity review for the following role(s) where "
            f"salary spread exceeds {int(PAY_EQUITY_THRESHOLD * 100)}%: {roles_flagged}. "
            f"Document and address any disparities to mitigate legal and retention risk."
        )

    # Recommendation 3 — missing performance reviews
    missing_count = len(turnover["missing_review"])
    if missing_count > 0:
        recs.append(
            f"Schedule performance reviews for the {missing_count} employee(s) with no "
            f"recorded review date. Timely reviews support goal alignment, development planning, "
            f"and defensible documentation in the event of disciplinary action."
        )

    # Recommendation 4 — high new-hire concentration in a department
    if not turnover["high_risk_depts"].empty:
        dept_names = ", ".join(turnover["high_risk_depts"]["department"].tolist())
        recs.append(
            f"The following department(s) have a high concentration of new hires relative "
            f"to total headcount: {dept_names}. Consider assigning formal mentors or buddy "
            f"partners to accelerate integration and reduce early attrition."
        )

    # Recommendation 5 — part-time workforce utilization
    pt_count = (headcount["ft_pt_by_dept"].get("Part-Time", pd.Series([0])).sum()
                if "Part-Time" in headcount["ft_pt_by_dept"].columns else 0)
    total = headcount["total_employees"]
    if pt_count > 0 and (pt_count / total) >= 0.15:
        recs.append(
            f"Part-time employees represent {pt_count} of {total} total headcount "
            f"({pt_count / total:.0%}). Review benefits eligibility, prorated PTO accrual "
            f"accuracy, and FLSA classification for all part-time roles to ensure ongoing "
            f"compliance."
        )

    # Fallback recommendation if data is clean
    if not recs:
        recs.append(
            "Workforce data appears healthy across all measured dimensions. Consider scheduling "
            "a mid-year engagement survey to proactively surface any employee concerns before "
            "they affect retention or performance."
        )

    # Tenure diversity recommendation (always include if we have enough data)
    avg_tenures = headcount["avg_tenure_by_dept"]["avg_tenure_years"]
    if not avg_tenures.empty and avg_tenures.max() - avg_tenures.min() > 2:
        low_dept = headcount["avg_tenure_by_dept"].loc[
            headcount["avg_tenure_by_dept"]["avg_tenure_years"].idxmin(), "department"
        ]
        recs.append(
            f"Tenure varies significantly across departments. The {low_dept} department has "
            f"the lowest average tenure, suggesting it may benefit from targeted retention "
            f"strategies such as career pathing conversations, skills development budgets, "
            f"or compensation benchmarking against market data."
        )

    return recs[:5]   # cap at 5


# ---------------------------------------------------------------------------
# Word document helpers (mirrors policy_generator.py style)
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_paragraph(doc, text="", bold=False, italic=False, font_size=11,
                  color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        run.font.color.rgb = color or COLOR_TEXT
    return p


def add_section_heading(doc, text: str):
    """Teal underlined section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEAL
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "007A87")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_body(doc, text: str, space_after=6, bold=False, italic=False, color=None):
    return add_paragraph(doc, text, font_size=10.5, space_after=space_after,
                         bold=bold, italic=italic, color=color)


def add_bullet(doc, text: str, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color or COLOR_TEXT
    return p


def fmt_currency(value) -> str:
    try:
        return f"${float(value):,.0f}"
    except (TypeError, ValueError):
        return "N/A"


def fmt_years(value) -> str:
    try:
        v = float(value)
        years  = int(v)
        months = round((v - years) * 12)
        if years == 0:
            return f"{months}mo"
        if months == 0:
            return f"{years}yr"
        return f"{years}yr {months}mo"
    except (TypeError, ValueError):
        return "N/A"


# ---------------------------------------------------------------------------
# Document structure helpers
# ---------------------------------------------------------------------------

def _add_table_header_row(table, headers: list[str], bg_hex: str = "1A3A5C"):
    """Style the first row of a table as a navy header row."""
    row = table.rows[0]
    for i, header in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, bg_hex)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_WHITE


def _add_data_row(table, row_idx: int, values: list, flag: bool = False):
    """Populate a data row; flag=True renders text in red."""
    row = table.rows[row_idx]
    text_color = COLOR_FLAG if flag else COLOR_TEXT
    bg = "FEF9F9" if flag else ("F4F6F8" if row_idx % 2 == 0 else "FFFFFF")
    for i, val in enumerate(values):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(str(val))
        run.font.size = Pt(9.5)
        run.font.color.rgb = text_color


def _make_table(doc, rows: int, cols: int) -> object:
    table = doc.add_table(rows=rows, cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    return table


# ---------------------------------------------------------------------------
# Report sections
# ---------------------------------------------------------------------------

def build_header(doc, company_name: str, report_date: str):
    if LOGO_PATH.exists():
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.paragraph_format.space_before = Pt(0)
        logo_para.paragraph_format.space_after  = Pt(8)
        logo_para.add_run().add_picture(str(LOGO_PATH), width=Inches(3.0))

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1A3A5C")

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(12)
    p1.paragraph_format.space_after  = Pt(2)
    r1 = p1.add_run("WORKFORCE ANALYSIS REPORT")
    r1.bold = True
    r1.font.size = Pt(16)
    r1.font.color.rgb = COLOR_WHITE

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(2)
    r2 = p2.add_run(company_name)
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0xB0, 0xC8, 0xD8)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after  = Pt(12)
    r3 = p3.add_run(
        f"Report Date: {report_date}  |  Prepared by {CONSULTANT_NAME}"
    )
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0xC8, 0xD8, 0xE8)

    doc.add_paragraph()


def build_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False

    ft = footer.paragraphs[0]
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ft.paragraph_format.space_before = Pt(4)

    pPr = ft._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "007A87")
    pBdr.append(top)
    pPr.append(pBdr)

    run = ft.add_run(
        f"{CONSULTANT_NAME}  |  {CONSULTANT_PHONE}  |  "
        f"{CONSULTANT_EMAIL}  |  {CONSULTANT_WEBSITE}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_MUTED

    p2 = footer.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(0)
    r_page = p2.add_run("Page ")
    r_page.font.size = Pt(8)
    r_page.font.color.rgb = COLOR_MUTED

    fldChar1  = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = "PAGE"
    fldChar2  = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run2 = p2.add_run()
    run2.font.size = Pt(8)
    run2.font.color.rgb = COLOR_MUTED
    run2._r.append(fldChar1)
    run2._r.append(instrText)
    run2._r.append(fldChar2)


def build_executive_summary(doc, headcount: dict, compensation: dict,
                             turnover: dict, company_name: str):
    add_section_heading(doc, "Executive Summary")
    total   = headcount["total_employees"]
    n_depts = len(headcount["total_by_dept"])
    avg_sal = compensation["avg_by_dept"]["avg_salary"].mean()
    n_flags = len(compensation["equity_flags"])
    n_new   = len(turnover["new_hires"])
    n_miss  = len(turnover["missing_review"])

    summary = (
        f"This Workforce Analysis Report presents a comprehensive review of "
        f"{company_name}'s workforce as of {date.today().strftime('%B %d, %Y')}. "
        f"The analysis covers {total} employees across {n_depts} department(s), "
        f"with an organization-wide average salary of {fmt_currency(avg_sal)}. "
    )
    if n_flags:
        summary += (
            f"Pay equity concerns were identified in {n_flags} role(s) and are "
            f"detailed in Section 2. "
        )
    if n_new:
        summary += (
            f"{n_new} employee(s) have been with the organization fewer than six months "
            f"and represent elevated onboarding and early-attrition risk. "
        )
    if n_miss:
        summary += (
            f"{n_miss} employee(s) have no recorded performance review date and require "
            f"immediate scheduling. "
        )
    summary += (
        "Recommended action items based on these findings are provided in Section 4."
    )
    add_body(doc, summary, space_after=10)


def build_headcount_section(doc, headcount: dict):
    add_section_heading(doc, "Section 1 — Headcount Summary")

    # --- 1a: Employees by department ---
    add_body(doc, "Total Employees by Department", bold=True, space_after=4)
    total_by_dept = headcount["total_by_dept"]
    table = _make_table(doc, rows=len(total_by_dept) + 1, cols=2)
    _add_table_header_row(table, ["Department", "Total Employees"])
    for i, row in total_by_dept.iterrows():
        _add_data_row(table, i + 1, [row["department"], int(row["total"])])
    doc.add_paragraph()

    # --- 1b: Full-Time vs Part-Time by department ---
    add_body(doc, "Full-Time vs. Part-Time Breakdown", bold=True, space_after=4)
    ft_pt = headcount["ft_pt_by_dept"].copy()
    ft_col = "Full-Time"  if "Full-Time"  in ft_pt.columns else None
    pt_col = "Part-Time"  if "Part-Time"  in ft_pt.columns else None
    headers = ["Department"]
    if ft_col:
        headers.append("Full-Time")
    if pt_col:
        headers.append("Part-Time")
    headers.append("Total")

    table = _make_table(doc, rows=len(ft_pt) + 1, cols=len(headers))
    _add_table_header_row(table, headers)
    for i, row in ft_pt.iterrows():
        ft_val = int(row[ft_col]) if ft_col else 0
        pt_val = int(row[pt_col]) if pt_col else 0
        values = [row["department"]]
        if ft_col:
            values.append(ft_val)
        if pt_col:
            values.append(pt_val)
        values.append(ft_val + pt_val)
        _add_data_row(table, i + 1, values)
    doc.add_paragraph()

    # --- 1c: Average tenure by department ---
    add_body(doc, "Average Tenure by Department", bold=True, space_after=4)
    tenure = headcount["avg_tenure_by_dept"]
    table = _make_table(doc, rows=len(tenure) + 1, cols=2)
    _add_table_header_row(table, ["Department", "Avg. Tenure"])
    for i, row in tenure.iterrows():
        _add_data_row(table, i + 1, [row["department"], fmt_years(row["avg_tenure_years"])])
    doc.add_paragraph()


def build_compensation_section(doc, compensation: dict):
    add_section_heading(doc, "Section 2 — Compensation Analysis")

    # --- 2a: Average salary by department ---
    add_body(doc, "Average Annual Salary by Department", bold=True, space_after=4)
    avg = compensation["avg_by_dept"]
    table = _make_table(doc, rows=len(avg) + 1, cols=2)
    _add_table_header_row(table, ["Department", "Avg. Annual Salary"])
    for i, row in avg.iterrows():
        _add_data_row(table, i + 1, [row["department"], fmt_currency(row["avg_salary"])])
    doc.add_paragraph()

    # --- 2b: Salary range per role ---
    add_body(doc, "Salary Range by Job Title", bold=True, space_after=4)
    by_role = compensation["salary_by_role"]
    table = _make_table(doc, rows=len(by_role) + 1, cols=5)
    _add_table_header_row(
        table,
        ["Job Title", "Employees", "Min Salary", "Max Salary", "Avg Salary"]
    )
    flagged_titles = {f["job_title"] for f in compensation["equity_flags"]}
    for i, row in by_role.iterrows():
        flag = row["job_title"] in flagged_titles
        _add_data_row(
            table, i + 1,
            [
                row["job_title"],
                int(row["count"]),
                fmt_currency(row["min_salary"]),
                fmt_currency(row["max_salary"]),
                fmt_currency(row["avg_salary"]),
            ],
            flag=flag,
        )
    doc.add_paragraph()

    # --- 2c: Pay equity flags ---
    if compensation["equity_flags"]:
        add_body(
            doc,
            "Pay Equity Concerns",
            bold=True, space_after=4, color=COLOR_FLAG
        )
        add_body(
            doc,
            f"The following role(s) contain two or more employees with a salary spread "
            f"exceeding {int(PAY_EQUITY_THRESHOLD * 100)}%. This may indicate inequitable "
            f"compensation practices and warrants a formal review. These rows are also "
            f"highlighted in red in the table above.",
            space_after=6, italic=True, color=COLOR_FLAG
        )
        for flag in compensation["equity_flags"]:
            add_bullet(
                doc,
                f"{flag['job_title']}: {flag['count']} employees — "
                f"range {fmt_currency(flag['min_salary'])} to "
                f"{fmt_currency(flag['max_salary'])} "
                f"({flag['spread_pct']:.1f}% spread)",
                color=COLOR_FLAG,
            )
        add_body(
            doc,
            "Recommendation: Engage qualified employment counsel before adjusting "
            "compensation for any flagged role, particularly where protected characteristics "
            "may intersect with pay disparities (FLSA / EPA compliance risk).",
            space_after=6, italic=True, color=COLOR_MUTED
        )
        doc.add_paragraph()
    else:
        add_body(
            doc,
            "No pay equity concerns identified. All roles with multiple incumbents "
            f"have salary spreads within the {int(PAY_EQUITY_THRESHOLD * 100)}% threshold.",
            space_after=6, italic=True
        )


def build_turnover_section(doc, turnover: dict):
    add_section_heading(doc, "Section 3 — Turnover Risk")

    # --- 3a: New hires (< 6 months) ---
    new_hires = turnover["new_hires"]
    add_body(
        doc,
        f"Employees with Tenure Under {NEW_HIRE_MONTHS} Months",
        bold=True, space_after=4
    )

    if new_hires.empty:
        add_body(doc, "No employees identified with tenure under six months.", italic=True)
    else:
        cols = ["full_name", "department", "job_title", "hire_date", "tenure_months"]
        display_cols = [c for c in cols if c in new_hires.columns]
        headers = {
            "full_name":      "Employee",
            "department":     "Department",
            "job_title":      "Job Title",
            "hire_date":      "Hire Date",
            "tenure_months":  "Months Tenure",
        }
        col_headers = [headers.get(c, c) for c in display_cols]
        table = _make_table(doc, rows=len(new_hires) + 1, cols=len(col_headers))
        _add_table_header_row(table, col_headers)
        for i, (_, row) in enumerate(new_hires.iterrows()):
            values = []
            for c in display_cols:
                if c == "hire_date":
                    values.append(
                        row[c].strftime("%Y-%m-%d") if pd.notna(row[c]) else "—"
                    )
                else:
                    values.append(row[c] if pd.notna(row[c]) else "—")
            _add_data_row(table, i + 1, values, flag=True)
    doc.add_paragraph()

    # --- 3b: High new-hire concentration departments ---
    add_body(
        doc,
        "Departments with High New-Hire Concentration",
        bold=True, space_after=4
    )
    high_risk = turnover["high_risk_depts"]
    if high_risk.empty:
        add_body(
            doc,
            "No departments identified with disproportionately high new-hire concentration.",
            italic=True
        )
    else:
        table = _make_table(doc, rows=len(high_risk) + 1, cols=4)
        _add_table_header_row(
            table, ["Department", "New Hires (<6 Mo)", "Total Headcount", "New-Hire Ratio"]
        )
        for i, row in high_risk.iterrows():
            _add_data_row(
                table, i + 1,
                [
                    row["department"],
                    int(row["new_hire_count"]),
                    int(row["total"]),
                    f"{row['new_hire_ratio']:.0%}",
                ],
                flag=True,
            )
    doc.add_paragraph()

    # --- 3c: Missing performance review dates ---
    add_body(doc, "Employees Missing Performance Review Date", bold=True, space_after=4)
    missing = turnover["missing_review"]
    if "last_performance_review" not in missing.columns or missing.empty:
        add_body(doc, "All employees have a recorded performance review date.", italic=True)
    else:
        cols = ["full_name", "department", "job_title", "hire_date"]
        display_cols = [c for c in cols if c in missing.columns]
        headers = {
            "full_name":  "Employee",
            "department": "Department",
            "job_title":  "Job Title",
            "hire_date":  "Hire Date",
        }
        col_headers = [headers.get(c, c) for c in display_cols]
        table = _make_table(doc, rows=len(missing) + 1, cols=len(col_headers))
        _add_table_header_row(table, col_headers)
        for i, (_, row) in enumerate(missing.iterrows()):
            values = []
            for c in display_cols:
                if c == "hire_date":
                    values.append(
                        row[c].strftime("%Y-%m-%d") if pd.notna(row[c]) else "—"
                    )
                else:
                    values.append(row[c] if pd.notna(row[c]) else "—")
            _add_data_row(table, i + 1, values, flag=True)
    doc.add_paragraph()


def build_recommendations_section(doc, recommendations: list[str]):
    add_section_heading(doc, "Section 4 — HR Recommendations")
    add_body(
        doc,
        "Based on the workforce data analyzed in this report, the following action items "
        "are recommended for consideration. Priority should be assigned in consultation with "
        "department leadership and, where applicable, qualified legal counsel.",
        space_after=8
    )
    for rec in recommendations:
        add_bullet(doc, rec)
    doc.add_paragraph()
    add_body(
        doc,
        "This report was prepared by Waymark HR Group LLC and is intended for internal "
        "HR and executive use. It does not constitute legal advice. Content involving "
        "compensation equity, leave policies, or termination procedures should be reviewed "
        "by qualified employment counsel before action is taken.",
        space_after=6, italic=True, color=COLOR_MUTED
    )


# ---------------------------------------------------------------------------
# Document assembly
# ---------------------------------------------------------------------------

def configure_document(doc):
    section = doc.sections[0]
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.85)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name      = "Calibri"
    style.font.size      = Pt(10.5)
    style.font.color.rgb = COLOR_TEXT


def generate_report(df: pd.DataFrame, company_name: str) -> Path:
    report_date = date.today().strftime("%B %d, %Y")

    # Run analysis
    headcount    = analyze_headcount(df)
    compensation = analyze_compensation(df)
    turnover     = analyze_turnover_risk(df)
    recs         = generate_recommendations(headcount, compensation, turnover)

    # Build document
    doc = Document()
    configure_document(doc)
    build_footer(doc)
    build_header(doc, company_name, report_date)
    build_executive_summary(doc, headcount, compensation, turnover, company_name)
    build_headcount_section(doc, headcount)
    build_compensation_section(doc, compensation)
    build_turnover_section(doc, turnover)
    build_recommendations_section(doc, recs)

    # Save
    name_slug  = company_name.replace(" ", "")
    date_slug  = date.today().strftime("%Y-%m-%d")
    filename   = f"{name_slug}_Workforce_Analysis_{date_slug}.docx"

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = REPORTS_DIR / filename
    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args():
    parser = argparse.ArgumentParser(
        description="Generate a Workforce Analysis Report from HR data.",
        epilog=(
            'Example: python workforce_analyzer.py '
            '--input data/spreadsheets/sample_workforce_data.csv '
            '--company "Acme Corporation"'
        )
    )
    parser.add_argument(
        "--input", "-i",
        required=True,
        help="Path to a CSV or Excel (.xlsx/.xls) workforce data file"
    )
    parser.add_argument(
        "--company", "-c",
        required=True,
        help="Client company name (used in the report title and filename)"
    )
    return parser.parse_args()


def main():
    args = parse_args()

    print("\n" + "=" * 60)
    print(f"  {CONSULTANT_NAME}")
    print("  Workforce Analyzer")
    print("=" * 60)
    print(f"  Company : {args.company}")
    print(f"  Input   : {args.input}")
    print("=" * 60)
    print("  Loading data...")

    df = load_data(args.input)

    print(f"  Loaded {len(df)} employee record(s) from {len(df['department'].unique())} department(s).")
    print("  Running analysis...")

    output_path = generate_report(df, args.company)

    print("=" * 60)
    print("  Workforce Analysis Report generated successfully.")
    print(f"  Saved to: {output_path}")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    main()
