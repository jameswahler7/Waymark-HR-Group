"""
onboarding_generator.py
Waymark HR Group LLC

Generates a professional employee onboarding welcome package as a Word document (.docx).
Saves the document to docs/reports/.

Usage:
    Interactive (prompts for each field):
        python onboarding_generator.py

    Non-interactive (all arguments provided):
        python onboarding_generator.py "Full Name" "Job Title" "Start Date" "Department" "Manager Name"

    Example:
        python onboarding_generator.py "Jane Smith" "HR Analyst" "March 24, 2026" "Human Resources" "Michael Johnson"
"""

import os
import sys
from datetime import datetime, date
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install it with: pip install python-docx")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

COMPANY_NAME = "Waymark HR Group LLC"
COMPANY_ADDRESS = "123 Corporate Drive, Suite 400 | Chicago, IL 60601"
COMPANY_PHONE = "716-225-6347"
COMPANY_EMAIL = "Jamie.Wahler@waymarkhrgroup.com"
COMPANY_WEBSITE = "www.waymarkhrgroup.com"

# Brand colors
COLOR_NAVY = RGBColor(0x1A, 0x3A, 0x5C)      # Dark navy
COLOR_TEAL = RGBColor(0x00, 0x7A, 0x87)       # Teal accent
COLOR_LIGHT_GRAY = RGBColor(0xF4, 0xF6, 0xF8) # Section background
COLOR_TEXT = RGBColor(0x2C, 0x2C, 0x2C)       # Body text
COLOR_WHITE = RGBColor(0xFF, 0xFF, 0xFF)

REPORTS_DIR = Path(__file__).resolve().parents[2] / "docs" / "reports"
LOGO_PATH   = Path(__file__).resolve().parents[2] / "docs" / "templates" / "Waymark_HR_Group_Logo_Full.png"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    """Set a table cell's background color via XML."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_paragraph(doc, text="", bold=False, italic=False, font_size=11,
                  color=None, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(font_size)
        run.font.color.rgb = color or COLOR_TEXT
    return p


def add_heading(doc, text, level=1):
    """Add a styled section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_TEAL
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "007A87")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_bullet(doc, text, indent_level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = COLOR_TEXT
    return p


def add_page_break(doc):
    doc.add_page_break()


def section_divider(doc):
    add_paragraph(doc, space_before=4, space_after=4)


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def build_header(doc, employee_name: str, job_title: str, start_date: str):
    """Full-width branded header block with company logo."""
    # Logo above the navy header
    if LOGO_PATH.exists():
        logo_para = doc.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_para.paragraph_format.space_before = Pt(0)
        logo_para.paragraph_format.space_after = Pt(8)
        logo_run = logo_para.add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Inches(3.0))

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_bg(cell, "1A3A5C")
    cell.width = Inches(6.5)

    p2 = cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(10)
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run("Employee Onboarding Welcome Package")
    run2.font.size = Pt(11)
    run2.font.color.rgb = RGBColor(0xB0, 0xC8, 0xD8)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(0)
    p3.paragraph_format.space_after = Pt(14)
    run3 = p3.add_run(f"{employee_name}  |  {job_title}  |  Start Date: {start_date}")
    run3.font.size = Pt(9.5)
    run3.font.color.rgb = RGBColor(0xC8, 0xD8, 0xE8)

    doc.add_paragraph()  # spacer


def build_welcome_letter(doc, employee_name: str, job_title: str,
                          start_date: str, department: str, manager_name: str):
    add_heading(doc, "Welcome Letter")

    add_paragraph(doc, datetime.now().strftime("%B %d, %Y"),
                  font_size=10.5, space_after=10)

    add_paragraph(doc, f"Dear {employee_name},", bold=True, font_size=11, space_after=8)

    body_paragraphs = [
        (f"On behalf of everyone at {COMPANY_NAME}, we are delighted to welcome you to "
         f"the team. Your appointment as {job_title} within our {department} department "
         f"reflects our confidence in the skills, expertise, and perspective you bring to "
         f"our organization."),
        ("Your first day will mark the beginning of an exciting professional journey. We "
         "have designed your onboarding experience to ensure you feel supported, informed, "
         "and empowered from the moment you arrive. This welcome package is your guide to "
         "everything you need to get started."),
        (f"Your direct manager, {manager_name}, will be in touch prior to your start date "
         "to confirm logistics and answer any preliminary questions you may have. Please "
         "do not hesitate to reach out to the HR team at any time — we are here to help."),
        ("We look forward to seeing the contributions you will make and are excited to have "
         "you as part of our team. Welcome aboard."),
    ]

    for para in body_paragraphs:
        add_paragraph(doc, para, font_size=10.5, space_after=8)

    add_paragraph(doc, "Warm regards,", font_size=10.5, space_after=2)
    add_paragraph(doc, "Human Resources Team", bold=True, font_size=10.5, space_after=2)
    add_paragraph(doc, COMPANY_NAME, font_size=10.5, space_after=2)
    add_paragraph(doc, COMPANY_PHONE, font_size=10, color=COLOR_TEAL, space_after=2)
    add_paragraph(doc, COMPANY_EMAIL, font_size=10, color=COLOR_TEAL, space_after=2)


def build_first_week_schedule(doc, start_date: str, manager_name: str, department: str):
    add_page_break(doc)
    add_heading(doc, "First Week Schedule")

    add_paragraph(
        doc,
        "The following schedule outlines planned activities for your first week. "
        "Your manager may adjust specific times based on team availability.",
        font_size=10, italic=True, space_after=10
    )

    days = [
        {
            "day": f"Day 1 — Monday ({start_date})",
            "activities": [
                "9:00 AM  —  Arrival, building access, and ID badge issuance",
                "9:30 AM  —  HR orientation: paperwork, benefits enrollment, and policies review",
                "11:00 AM —  IT setup: workstation, email, system access provisioning",
                f"12:00 PM —  Welcome lunch with {manager_name} and immediate team",
                "1:30 PM  —  Office tour and introduction to key stakeholders",
                "3:00 PM  —  Review role expectations and 30/60/90-day goals",
                "4:30 PM  —  Independent review of provided onboarding materials",
            ],
        },
        {
            "day": "Day 2 — Tuesday",
            "activities": [
                f"9:00 AM  —  Team meeting with {department} department",
                "10:00 AM —  Systems and tools training (core platforms)",
                "12:00 PM —  Lunch (self-directed)",
                "1:00 PM  —  Introduction to current projects and priorities",
                "3:00 PM  —  1:1 check-in with manager",
                "4:00 PM  —  Independent review: internal documentation and workflows",
            ],
        },
        {
            "day": "Day 3 — Wednesday",
            "activities": [
                "9:00 AM  —  Cross-department introductions (scheduled by manager)",
                "11:00 AM —  Compliance training: harassment prevention, data privacy",
                "12:00 PM —  Lunch (self-directed)",
                "1:00 PM  —  Shadow a colleague on active project or client work",
                "3:30 PM  —  Benefits orientation follow-up and enrollment deadline review",
            ],
        },
        {
            "day": "Day 4 — Thursday",
            "activities": [
                "9:00 AM  —  Continue systems and tools training",
                "11:00 AM —  Review company policies and employee handbook",
                "12:00 PM —  Lunch (self-directed)",
                "1:00 PM  —  Begin first assigned task or project contribution",
                "4:00 PM  —  End-of-day check-in with manager",
            ],
        },
        {
            "day": "Day 5 — Friday",
            "activities": [
                "9:00 AM  —  Independent work on assigned tasks",
                "11:00 AM —  End-of-week reflection: questions, feedback, and open discussion",
                "12:00 PM —  Team lunch (optional, as available)",
                "1:30 PM  —  Complete any outstanding paperwork or system enrollments",
                "3:00 PM  —  Week-one recap meeting with manager",
                "4:30 PM  —  Set goals and priorities for Week 2",
            ],
        },
    ]

    for day_block in days:
        add_paragraph(doc, day_block["day"], bold=True, font_size=10.5,
                      color=COLOR_NAVY, space_before=8, space_after=3)
        for activity in day_block["activities"]:
            add_bullet(doc, activity)

    section_divider(doc)


def build_hr_policies_summary(doc):
    add_page_break(doc)
    add_heading(doc, "Key HR Policies Summary")

    add_paragraph(
        doc,
        "The following is a high-level summary of key company policies. Employees are "
        "responsible for reading and adhering to all policies in the full Employee Handbook, "
        "which will be provided separately. Please direct any policy questions to the HR team.",
        font_size=10, italic=True, space_after=10
    )

    policies = [
        {
            "title": "Code of Conduct",
            "points": [
                "Employees are expected to conduct themselves with integrity, professionalism, "
                "and respect in all workplace interactions.",
                "Conflicts of interest must be disclosed promptly to a supervisor or HR.",
                "Confidential company and client information must be protected at all times.",
            ],
        },
        {
            "title": "Anti-Harassment and Equal Opportunity",
            "points": [
                "The company maintains a zero-tolerance policy for harassment, discrimination, "
                "or retaliation of any kind.",
                "All individuals are entitled to a work environment free from unlawful "
                "discrimination based on race, color, religion, sex, national origin, age, "
                "disability, or any other protected characteristic.",
                "Concerns should be reported to HR immediately. All reports are handled "
                "confidentially to the extent possible.",
            ],
        },
        {
            "title": "Attendance and Work Hours",
            "points": [
                "Core business hours are Monday through Friday, 8:00 AM – 5:00 PM local time, "
                "unless otherwise specified by your department.",
                "Remote or hybrid work arrangements must be approved in writing by your manager "
                "and HR.",
                "Absences must be reported to your manager as early as possible, and no later "
                "than the start of the scheduled workday.",
            ],
        },
        {
            "title": "Paid Time Off (PTO)",
            "points": [
                "PTO accrual begins on your first day of employment per the company's published "
                "accrual schedule.",
                "PTO requests should be submitted through the HR system at least two weeks in "
                "advance when possible.",
                "Unused PTO is subject to the company's carryover and payout policy as outlined "
                "in the Employee Handbook.",
            ],
        },
        {
            "title": "Performance Reviews",
            "points": [
                "Formal performance reviews are conducted annually, typically in Q4.",
                "New employees receive a 90-day check-in review within their first quarter.",
                "Performance feedback is an ongoing process; employees are encouraged to engage "
                "in regular dialogue with their manager.",
            ],
        },
        {
            "title": "Data Privacy and Confidentiality",
            "points": [
                "Employees must comply with all applicable data privacy regulations, including "
                "the handling of client and employee personal data.",
                "Company data must not be shared with unauthorized parties or stored on "
                "unmanaged personal devices.",
                "Any suspected data breach or unauthorized disclosure must be reported to IT "
                "and HR immediately.",
            ],
        },
    ]

    for policy in policies:
        add_paragraph(doc, policy["title"], bold=True, font_size=10.5,
                      color=COLOR_NAVY, space_before=10, space_after=3)
        for point in policy["points"]:
            add_bullet(doc, point)

    add_paragraph(
        doc,
        "Note: This summary does not constitute legal advice. Employees should consult the "
        "full Employee Handbook or speak with HR for complete policy details.",
        font_size=9.5, italic=True, color=RGBColor(0x88, 0x88, 0x88), space_before=12
    )


def build_it_checklist(doc, employee_name: str):
    add_page_break(doc)
    add_heading(doc, "IT Setup Checklist")

    add_paragraph(
        doc,
        "Please ensure the following items are completed during your first week. "
        "Contact the IT Help Desk for assistance with any item on this list.",
        font_size=10, italic=True, space_after=10
    )

    categories = [
        {
            "category": "Hardware & Workspace",
            "items": [
                "Workstation or laptop assigned and configured",
                "External monitor, keyboard, and mouse (if applicable)",
                "Phone or softphone setup completed",
                "VPN client installed and tested",
                "Security badge / physical access activated",
            ],
        },
        {
            "category": "Accounts & Access",
            "items": [
                "Company email account created and accessible",
                "Multi-factor authentication (MFA) enabled on all accounts",
                "Password manager account created (if company-issued)",
                "Core business applications provisioned (HR system, project tools, etc.)",
                "Shared drives and department file access confirmed",
                "Video conferencing platform (e.g., Zoom, Teams) installed and tested",
            ],
        },
        {
            "category": "Security & Compliance",
            "items": [
                "Acceptable Use Policy reviewed and signed",
                "Endpoint protection / antivirus confirmed active",
                "Device encryption verified with IT",
                "Security awareness training module completed",
                "IT security contact information saved",
            ],
        },
        {
            "category": "Communication & Collaboration",
            "items": [
                "Email signature configured per company template",
                "Instant messaging / chat platform installed (Slack, Teams, etc.)",
                "Calendar shared with manager and team (as applicable)",
                "Video conferencing profile photo and display name set",
                "Company intranet or knowledge base access confirmed",
            ],
        },
    ]

    for cat in categories:
        add_paragraph(doc, cat["category"], bold=True, font_size=10.5,
                      color=COLOR_NAVY, space_before=10, space_after=3)

        # Draw a simple checkbox table
        table = doc.add_table(rows=len(cat["items"]), cols=2)
        table.style = "Table Grid"
        col_widths = [Inches(0.4), Inches(5.6)]

        for i, item in enumerate(cat["items"]):
            row = table.rows[i]
            row.cells[0].width = col_widths[0]
            row.cells[1].width = col_widths[1]

            cb_para = row.cells[0].paragraphs[0]
            cb_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cb_run = cb_para.add_run("☐")
            cb_run.font.size = Pt(11)

            item_para = row.cells[1].paragraphs[0]
            item_para.paragraph_format.space_before = Pt(2)
            item_para.paragraph_format.space_after = Pt(2)
            item_run = item_para.add_run(item)
            item_run.font.size = Pt(10.5)
            item_run.font.color.rgb = COLOR_TEXT

        doc.add_paragraph()  # spacer after table

    add_paragraph(
        doc,
        f"IT Help Desk  |  helpdesk@waymarkhrgroup.com  |  Ext. 200",
        font_size=9.5, color=COLOR_TEAL, space_before=6
    )


def build_emergency_contacts(doc, manager_name: str):
    add_page_break(doc)
    add_heading(doc, "Emergency Contacts & Key Directory")

    add_paragraph(
        doc,
        "Please retain this information for quick reference. Emergency procedures and "
        "building evacuation plans are posted at all building exits.",
        font_size=10, italic=True, space_after=10
    )

    # Internal contacts table
    add_paragraph(doc, "Internal Contacts", bold=True, font_size=10.5,
                  color=COLOR_NAVY, space_after=4)

    internal_contacts = [
        ("Role", "Name", "Phone / Extension", "Email"),
        ("Your Direct Manager", manager_name, "[Manager Phone]", "[Manager Email]"),
        ("HR Department", "HR Team", COMPANY_PHONE + " Ext. 100", COMPANY_EMAIL),
        ("IT Help Desk", "IT Support", "Ext. 200", "helpdesk@waymarkhrgroup.com"),
        ("Facilities / Office Manager", "[Facilities Contact]", "Ext. 300",
         "facilities@waymarkhrgroup.com"),
        ("Payroll Department", "Payroll Team", "Ext. 105", "payroll@waymarkhrgroup.com"),
        ("Benefits Administrator", "Benefits Team", "Ext. 106", "benefits@waymarkhrgroup.com"),
    ]

    table = doc.add_table(rows=len(internal_contacts), cols=4)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(internal_contacts):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            run = para.add_run(cell_text)
            run.font.size = Pt(9.5)
            if row_idx == 0:
                run.bold = True
                run.font.color.rgb = COLOR_WHITE
                set_cell_bg(cell, "1A3A5C")
            else:
                run.font.color.rgb = COLOR_TEXT

    doc.add_paragraph()

    # Emergency services
    add_paragraph(doc, "Emergency Services", bold=True, font_size=10.5,
                  color=COLOR_NAVY, space_before=8, space_after=4)

    emergency_contacts = [
        ("Emergency", "Police / Fire / Ambulance", "911"),
        ("Non-Emergency Police", "Local Precinct", "[Local Phone Number]"),
        ("Poison Control", "National Hotline", "1-800-222-1222"),
        ("Mental Health Crisis Line", "988 Suicide & Crisis Lifeline", "988"),
        ("Employee Assistance Program (EAP)", "EAP Provider", "[EAP Phone Number]"),
    ]

    table2 = doc.add_table(rows=len(emergency_contacts) + 1, cols=3)
    table2.style = "Table Grid"

    # Header row
    headers = ["Type", "Contact", "Number"]
    header_row = table2.rows[0]
    for col_idx, header in enumerate(headers):
        cell = header_row.cells[col_idx]
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(3)
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(header)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_WHITE
        set_cell_bg(cell, "007A87")

    for row_idx, (etype, contact, number) in enumerate(emergency_contacts, start=1):
        row = table2.rows[row_idx]
        for col_idx, cell_text in enumerate([etype, contact, number]):
            cell = row.cells[col_idx]
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            run = para.add_run(cell_text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = COLOR_TEXT

    doc.add_paragraph()

    # Employee personal emergency contact block
    add_paragraph(doc, "Your Personal Emergency Contact Information", bold=True,
                  font_size=10.5, color=COLOR_NAVY, space_before=12, space_after=4)
    add_paragraph(
        doc,
        "Please provide this information to HR for our records. It will only be used "
        "in the event of a workplace emergency.",
        font_size=10, italic=True, space_after=8
    )

    personal_fields = [
        ("Emergency Contact Name:", ""),
        ("Relationship:", ""),
        ("Primary Phone:", ""),
        ("Secondary Phone:", ""),
        ("Address (optional):", ""),
    ]

    for label, _ in personal_fields:
        table3 = doc.add_table(rows=1, cols=2)
        table3.style = "Table Grid"
        row = table3.rows[0]
        row.cells[0].width = Inches(2.2)
        row.cells[1].width = Inches(4.3)
        lbl_para = row.cells[0].paragraphs[0]
        lbl_para.paragraph_format.space_before = Pt(4)
        lbl_para.paragraph_format.space_after = Pt(4)
        lbl_run = lbl_para.add_run(label)
        lbl_run.bold = True
        lbl_run.font.size = Pt(9.5)
        lbl_run.font.color.rgb = COLOR_TEXT

        val_para = row.cells[1].paragraphs[0]
        val_para.paragraph_format.space_before = Pt(4)
        val_para.paragraph_format.space_after = Pt(4)
        val_run = val_para.add_run(" ")
        val_run.font.size = Pt(9.5)

        doc.add_paragraph()


def build_footer_note(doc):
    """Add a closing footer note."""
    add_paragraph(doc, space_before=12, space_after=4)
    # Horizontal rule via paragraph border
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "4")
    top.set(qn("w:space"), "1")
    top.set(qn("w:color"), "007A87")
    pBdr.append(top)
    pPr.append(pBdr)

    footer_lines = [
        f"{COMPANY_NAME}  |  {COMPANY_ADDRESS}",
        f"{COMPANY_PHONE}  |  {COMPANY_EMAIL}  |  {COMPANY_WEBSITE}",
        "This document is confidential and intended solely for the named recipient. "
        "Contents are subject to change. Please contact HR with any questions.",
    ]
    for line in footer_lines:
        add_paragraph(doc, line, font_size=8.5, color=RGBColor(0x88, 0x88, 0x88),
                      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def configure_document(doc):
    """Set page margins and default font."""
    from docx.oxml.ns import nsmap
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = COLOR_TEXT


# ---------------------------------------------------------------------------
# Input collection
# ---------------------------------------------------------------------------

def prompt_employee_details() -> dict:
    print("\n" + "=" * 60)
    print(f"  {COMPANY_NAME}")
    print("  Onboarding Welcome Package Generator")
    print("=" * 60 + "\n")
    print("Please enter the following employee details:\n")

    fields = [
        ("employee_name", "Full Name"),
        ("job_title",     "Job Title"),
        ("start_date",    "Start Date (e.g., March 24, 2026)"),
        ("department",    "Department"),
        ("manager_name",  "Manager's Full Name"),
    ]

    details = {}
    for key, label in fields:
        while True:
            value = input(f"  {label}: ").strip()
            if value:
                details[key] = value
                break
            print(f"    ⚠  {label} cannot be blank. Please try again.")

    print()
    return details


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def generate_onboarding_package(details: dict) -> Path:
    doc = Document()
    configure_document(doc)

    build_header(doc, details["employee_name"], details["job_title"], details["start_date"])
    build_welcome_letter(doc, details["employee_name"], details["job_title"],
                         details["start_date"], details["department"], details["manager_name"])
    build_first_week_schedule(doc, details["start_date"],
                               details["manager_name"], details["department"])
    build_hr_policies_summary(doc)
    build_it_checklist(doc, details["employee_name"])
    build_emergency_contacts(doc, details["manager_name"])
    build_footer_note(doc)

    # Filename: LastFirst_Onboarding_YYYY-MM-DD.docx
    name_slug = details["employee_name"].replace(" ", "")
    date_slug = datetime.now().strftime("%Y-%m-%d")
    filename = f"{name_slug}_Onboarding_{date_slug}.docx"

    REPORTS_DIR.mkdir(parents=True, exist_ok=True)
    output_path = REPORTS_DIR / filename
    doc.save(output_path)
    return output_path


def parse_args():
    import argparse
    parser = argparse.ArgumentParser(
        description="Generate an employee onboarding welcome package.",
        epilog="If no arguments are provided, the script will prompt for each value interactively."
    )
    parser.add_argument("full_name",    nargs="?", help="Employee's full name")
    parser.add_argument("job_title",    nargs="?", help="Employee's job title")
    parser.add_argument("start_date",   nargs="?", help='Start date (e.g., "March 24, 2026")')
    parser.add_argument("department",   nargs="?", help="Employee's department")
    parser.add_argument("manager_name", nargs="?", help="Manager's full name")
    return parser.parse_args()


def main():
    args = parse_args()

    all_provided = all([args.full_name, args.job_title, args.start_date,
                        args.department, args.manager_name])

    if all_provided:
        details = {
            "employee_name": args.full_name,
            "job_title":     args.job_title,
            "start_date":    args.start_date,
            "department":    args.department,
            "manager_name":  args.manager_name,
        }
    else:
        details = prompt_employee_details()

    print("  Generating onboarding package...")
    output_path = generate_onboarding_package(details)

    print("\n" + "=" * 60)
    print("  Onboarding package generated successfully.")
    print(f"  Saved to: {output_path}")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    main()
